"""
운행일지 내보내기: Excel / PDF / DOCX (한컴 호환)
⚠️ 워터마크 절대 금지 — Free/Pro 모두 완전 클린 문서 (회사 제출 가능)
"""

from __future__ import annotations

import io
from datetime import datetime
from typing import Any

import pandas as pd

from modules.validator import format_minutes_kr


# ────────────────────────────────────────────────────────
# 공통 헬퍼
# ────────────────────────────────────────────────────────


# 기본 서식 팔레트 — 회사 제출·유튜브 시연용 (서비스 브랜드 노출 없음)
_FORM_NAVY = "0F172A"
_FORM_SLATE = "334155"
_FORM_LINE = "94A3B8"
_FORM_GRID = "CBD5E1"
_FORM_HEAD_BG = "1E293B"
_FORM_ZEBRA = "F8FAFC"
_FORM_LABEL_BG = "F1F5F9"
_FORM_TOTAL_BG = "EEF2FF"
_FORM_TITLE = "차량 운행일지"
# A4 한 장을 서식이 꽉 채우도록 구간 표 최소 행 수 (데이터 부족 시 빈칸 행)
_FORM_MIN_TRIP_ROWS = 12
_FORM_MIN_VISIT_ROWS = 10


def _pad_trip_records(records: list[dict], min_rows: int = _FORM_MIN_TRIP_ROWS) -> list[dict]:
    """구간 데이터가 적어도 A4 양식 칸이 비지 않도록 빈 행으로 채움."""
    out = list(records or [])
    empty = {
        "순번": "",
        "출발시각": "",
        "도착시각": "",
        "출발지": "",
        "도착지": "",
        "운행목적": "",
        "거리(km)": "",
        "운행시간": "",
        "비고": "",
    }
    while len(out) < min_rows:
        blank = dict(empty)
        blank["순번"] = len(out) + 1
        out.append(blank)
    # 순번 재부여 (빈 행 포함)
    for i, rec in enumerate(out, start=1):
        if not rec.get("출발시각") and not rec.get("출발지") and not rec.get("도착지"):
            rec["순번"] = i
        else:
            rec["순번"] = i
    return out


def _trips_dataframe(log: dict) -> pd.DataFrame:
    """구간 표 — 회사 제출용. 점심제외(분) 칸은 넣지 않음(서식에 거의 없고 0만 반복)."""
    rows = []
    for i, t in enumerate(log.get("trips") or [], start=1):
        rows.append(
            {
                "순번": i,
                "출발시각": t.get("depart_time", ""),
                "도착시각": t.get("arrive_time", ""),
                "출발지": t.get("from", ""),
                "도착지": t.get("to", ""),
                "운행목적": t.get("purpose", ""),
                "거리(km)": t.get("distance_km", ""),
                "운행시간": t.get("duration_display", ""),
                "비고": t.get("memo", ""),
            }
        )
    return pd.DataFrame(rows)


def _fuel_label(meta: dict) -> str:
    if meta.get("주유 금액"):
        label = str(meta["주유 금액"])
    elif meta.get("주유"):
        label = str(meta["주유"])
    else:
        label = "—"
    if meta.get("주유량"):
        label = f"{label} · {meta['주유량']}"
    return label


def _meta_display_rows(log: dict, meta: dict | None = None) -> list[tuple[str, str, str, str]]:
    """메타 2열×N행 (라벨,값,라벨,값) — 인쇄·엑셀·PDF 공통. 라벨은 칸에 안 잘리게 짧게."""
    meta = meta or _meta(log)
    odo_s = meta.get("최초 누적(km)") or ""
    odo_e = meta.get("종료 누적(km)") or ""
    if odo_s or odo_e:
        odo_txt = f"{odo_s or '—'} → {odo_e or '—'} km"
    else:
        odo_txt = "—"
    dist = meta.get("총 거리(km)") or ""
    dist_txt = f"{dist} km" if dist not in ("", None) else "—"
    return [
        ("작성일", meta.get("작성일") or "—", "차량번호", meta.get("차량번호") or "—"),
        ("운전자", meta.get("운전자") or "—", "회사명", meta.get("회사명") or "—"),
        ("누적거리", odo_txt, "총 거리", dist_txt),
        ("운행시간", meta.get("총 운행시간") or "—", "주유", _fuel_label(meta)),
    ]


def _privacy_log(log: dict) -> dict:
    """내보내기 전 점심 장소(사생활) 스크럽. 설정 기본=비노출."""
    try:
        from modules.generator import scrub_lunch_place_privacy

        # 문서 제출물에는 기본 비노출. 명시적으로 lunch_place 가 있어도
        # omit 기본값으로 제거 (사용자 설정이 로그에 실린 경우 예외 처리 가능)
        settings = log.get("_settings") if isinstance(log.get("_settings"), dict) else {}
        if log.get("include_lunch_place") is True or settings.get("include_lunch_place") is True:
            settings = {**settings, "omit_lunch_place": False}
        else:
            settings = {**settings, "omit_lunch_place": settings.get("omit_lunch_place", True)}
        return scrub_lunch_place_privacy(log, settings)
    except Exception:
        out = dict(log or {})
        out["lunch_place"] = ""
        return out


def _meta(log: dict) -> dict[str, str]:
    log = _privacy_log(log)
    meta = {
        "작성일": str(log.get("date") or datetime.now().strftime("%Y-%m-%d")),
        "차량번호": str(log.get("vehicle") or ""),
        "운전자": str(log.get("driver_name") or ""),
        "회사명": str(log.get("company_name") or ""),
        "최초 누적(km)": str(log.get("odometer_start") if log.get("odometer_start") is not None else ""),
        "종료 누적(km)": str(log.get("odometer_end") if log.get("odometer_end") is not None else ""),
        "총 거리(km)": str(log.get("total_distance_km") or ""),
        "총 운행시간": format_minutes_kr(int(log.get("total_net_minutes") or 0)),
        # 점심 제외분은 내부 검증·순수운행 계산에만 사용 (제출 문서 칸으로는 노출하지 않음)
        "요약": str(log.get("summary") or ""),
    }
    if log.get("fuel_refueled"):
        amt = log.get("fuel_amount_krw")
        liters = log.get("fuel_liters")
        if amt is not None and amt != "":
            try:
                meta["주유 금액"] = f"{int(float(amt)):,}원"
            except (TypeError, ValueError):
                meta["주유 금액"] = f"{amt}원"
        else:
            meta["주유"] = "함"
        if liters is not None and liters != "":
            meta["주유량"] = f"{liters} L"
    else:
        meta["주유"] = "안 함"
    # 점심 장소는 사생활 — 비어 있지 않고 명시 허용된 경우만 포함
    lunch = str(log.get("lunch_place") or "").strip()
    if lunch and (
        log.get("include_lunch_place") is True
        or (isinstance(log.get("_settings"), dict) and log["_settings"].get("include_lunch_place") is True)
    ):
        meta["점심 장소"] = lunch
    return meta


def _filename_base(log: dict) -> str:
    d = str(log.get("date") or datetime.now().strftime("%Y-%m-%d")).replace("-", "")
    if str(log.get("report_type") or "").lower() in ("field", "field_visit", "outing"):
        who = str(log.get("author_name") or log.get("driver_name") or "field").replace(" ", "")
        return f"외근일지_{d}_{who or 'field'}"
    v = str(log.get("vehicle") or "vehicle").replace(" ", "")
    return f"운행일지_{d}_{v}"


def _is_field_log(log: dict) -> bool:
    return str(log.get("report_type") or "").lower() in ("field", "field_visit", "outing")


def _field_visits(log: dict) -> list[dict]:
    visits = log.get("visits")
    if isinstance(visits, list) and visits:
        return [v for v in visits if isinstance(v, dict)]
    # trips 호환
    out = []
    for t in log.get("trips") or []:
        if not isinstance(t, dict):
            continue
        out.append(
            {
                "time": t.get("depart_time") or t.get("time") or "",
                "place": t.get("to") or t.get("place") or "",
                "purpose": t.get("purpose") or "",
                "result": t.get("result") or "",
                "next_action": t.get("next_action") or "",
                "memo": t.get("memo") or "",
            }
        )
    return out


def _export_field_excel(log: dict) -> tuple[bytes, str]:
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
    from openpyxl.utils import get_column_letter

    wb = Workbook()
    ws = wb.active
    ws.title = "외근일지"
    ws.sheet_view.showGridLines = False
    thin = Border(
        left=Side(style="thin", color=_FORM_GRID),
        right=Side(style="thin", color=_FORM_GRID),
        top=Side(style="thin", color=_FORM_GRID),
        bottom=Side(style="thin", color=_FORM_GRID),
    )
    header_fill = PatternFill("solid", fgColor=_FORM_HEAD_BG)
    label_fill = PatternFill("solid", fgColor=_FORM_LABEL_BG)
    zebra_fill = PatternFill("solid", fgColor=_FORM_ZEBRA)
    white_fill = PatternFill("solid", fgColor="FFFFFF")
    header_font = Font(bold=True, color="FFFFFF", name="맑은 고딕", size=10)
    title_font = Font(bold=True, color="FFFFFF", name="맑은 고딕", size=16)
    sub_font = Font(color="CBD5E1", name="맑은 고딕", size=9)
    label_font = Font(bold=True, color=_FORM_SLATE, name="맑은 고딕", size=9)
    cell_font = Font(name="맑은 고딕", size=10, color=_FORM_NAVY)
    center = Alignment(horizontal="center", vertical="center", wrap_text=True)
    left = Alignment(horizontal="left", vertical="center", wrap_text=True)

    n_cols = 7
    ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=n_cols)
    c = ws.cell(row=1, column=1, value="외근·출장 일지")
    c.font = title_font
    c.fill = PatternFill("solid", fgColor=_FORM_NAVY)
    c.alignment = center
    ws.row_dimensions[1].height = 34

    ws.merge_cells(start_row=2, start_column=1, end_row=2, end_column=n_cols)
    s = ws.cell(row=2, column=1, value="방문·업무 결과 기록 · 제출용")
    s.font = sub_font
    s.fill = PatternFill("solid", fgColor=_FORM_HEAD_BG)
    s.alignment = center
    ws.row_dimensions[2].height = 18

    meta_pairs = [
        ("작성일", str(log.get("date") or "—"), "작성자", str(log.get("author_name") or log.get("driver_name") or "—")),
        ("부서", str(log.get("department") or "—"), "회사명", str(log.get("company_name") or "—")),
    ]
    r = 4
    for lab1, val1, lab2, val2 in meta_pairs:
        ws.cell(row=r, column=1, value=lab1).font = label_font
        ws.cell(row=r, column=1).fill = label_fill
        ws.cell(row=r, column=1).border = thin
        ws.cell(row=r, column=1).alignment = center
        ws.merge_cells(start_row=r, start_column=2, end_row=r, end_column=3)
        ws.cell(row=r, column=2, value=val1).font = cell_font
        for col in range(2, 4):
            ws.cell(row=r, column=col).border = thin
            ws.cell(row=r, column=col).fill = white_fill
        ws.cell(row=r, column=4, value=lab2).font = label_font
        ws.cell(row=r, column=4).fill = label_fill
        ws.cell(row=r, column=4).border = thin
        ws.cell(row=r, column=4).alignment = center
        ws.merge_cells(start_row=r, start_column=5, end_row=r, end_column=7)
        ws.cell(row=r, column=5, value=val2).font = cell_font
        for col in range(5, 8):
            ws.cell(row=r, column=col).border = thin
            ws.cell(row=r, column=col).fill = white_fill
        ws.row_dimensions[r].height = 22
        r += 1

    r += 1
    ws.cell(row=r, column=1, value="업무 요약").font = label_font
    ws.cell(row=r, column=1).fill = label_fill
    ws.cell(row=r, column=1).border = thin
    ws.cell(row=r, column=1).alignment = center
    ws.merge_cells(start_row=r, start_column=2, end_row=r, end_column=n_cols)
    ws.cell(row=r, column=2, value=str(log.get("summary") or "")).font = cell_font
    for col in range(2, n_cols + 1):
        ws.cell(row=r, column=col).border = thin
        ws.cell(row=r, column=col).fill = white_fill
    ws.row_dimensions[r].height = 28
    r += 2

    headers = ["순번", "시각", "방문처", "목적", "결과", "후속 조치", "비고"]
    header_row = r
    for i, h in enumerate(headers, start=1):
        cell = ws.cell(row=header_row, column=i, value=h)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = center
        cell.border = thin
    ws.row_dimensions[header_row].height = 24

    visits = _field_visits(log)
    for r_idx, v in enumerate(visits, start=1):
        vals = [
            r_idx,
            v.get("time") or "",
            v.get("place") or "",
            v.get("purpose") or "",
            v.get("result") or "",
            v.get("next_action") or "",
            v.get("memo") or "",
        ]
        fill = zebra_fill if r_idx % 2 == 0 else white_fill
        for c_idx, val in enumerate(vals, start=1):
            cell = ws.cell(row=header_row + r_idx, column=c_idx, value=val)
            cell.font = cell_font
            cell.border = thin
            cell.fill = fill
            cell.alignment = center if c_idx in (1, 2) else left

    sig_row = header_row + len(visits) + 2
    for label, col_s, col_e in (("작성자", 4, 5), ("확인자", 6, 7)):
        ws.merge_cells(start_row=sig_row, start_column=col_s, end_row=sig_row, end_column=col_e)
        ws.merge_cells(start_row=sig_row + 1, start_column=col_s, end_row=sig_row + 2, end_column=col_e)
        head = ws.cell(row=sig_row, column=col_s, value=label)
        head.font = label_font
        head.fill = label_fill
        head.alignment = center
        for col in range(col_s, col_e + 1):
            ws.cell(row=sig_row, column=col).border = thin
            ws.cell(row=sig_row, column=col).fill = label_fill
        body = ws.cell(row=sig_row + 1, column=col_s, value="(서명)")
        body.font = Font(name="맑은 고딕", size=9, color="94A3B8")
        body.alignment = center
        for rr in (sig_row + 1, sig_row + 2):
            for col in range(col_s, col_e + 1):
                ws.cell(row=rr, column=col).border = thin

    widths = [6, 10, 20, 14, 20, 16, 14]
    for i, w in enumerate(widths, start=1):
        ws.column_dimensions[get_column_letter(i)].width = w

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue(), f"{_filename_base(log)}.xlsx"


# ────────────────────────────────────────────────────────
# Excel (클린 — 워터마크 없음)
# ────────────────────────────────────────────────────────


def _default_driving_template_path() -> "Path":
    """기본 운행일지 엑셀 서식 경로 (칸 확대 A4 양식)."""
    from pathlib import Path

    root = Path(__file__).resolve().parents[1]
    candidates = [
        root / "web" / "assets" / "templates" / "default_driving_log.xlsx",
        root / "assets" / "forms" / "default_driving_log.xlsx",
    ]
    for p in candidates:
        if p.is_file():
            return p
    return candidates[0]


def export_excel(log: dict) -> tuple[bytes, str]:
    """
    기본 엑셀 내보내기.
    고정 서식 템플릿(default_driving_log.xlsx)에 데이터만 채움 —
    항상 칸 확대 A4 양식으로 동일하게 출력.
    """
    log = _privacy_log(log)
    if _is_field_log(log):
        return _export_field_excel(log)

    from openpyxl import load_workbook

    tmpl = _default_driving_template_path()
    if not tmpl.is_file():
        # 템플릿 누락 시 코드 생성 폴백
        return _export_excel_built(log)

    wb = load_workbook(tmpl)
    ws = wb.active

    meta = _meta(log)
    rows_meta = _meta_display_rows(log, meta)
    # 템플릿 레이아웃 고정:
    # A4/B4 작성일·값 | E4/F4 차량번호·값
    # A5/B5 운전자 | E5/F5 회사명
    # A6/B6 누적거리 | E6/F6 총 거리
    # A7/B7 운행시간 | E7/F7 주유
    # A9/B9 운행 요약
    # A10 헤더, A11–A22 구간 12행, A23 합계, E25/G25 서명
    meta_value_cells = [
        ("B4", rows_meta[0][1]),
        ("F4", rows_meta[0][3]),
        ("B5", rows_meta[1][1]),
        ("F5", rows_meta[1][3]),
        ("B6", rows_meta[2][1]),
        ("F6", rows_meta[2][3]),
        ("B7", rows_meta[3][1]),
        ("F7", rows_meta[3][3]),
    ]
    for coord, val in meta_value_cells:
        ws[coord] = val if val not in (None, "") else None

    # 라벨도 템플릿과 동기 (누적거리 등)
    ws["A4"] = rows_meta[0][0]
    ws["E4"] = rows_meta[0][2]
    ws["A5"] = rows_meta[1][0]
    ws["E5"] = rows_meta[1][2]
    ws["A6"] = rows_meta[2][0]
    ws["E6"] = rows_meta[2][2]
    ws["A7"] = rows_meta[3][0]
    ws["E7"] = rows_meta[3][2]
    ws["B9"] = meta.get("요약") or None

    trips = list(log.get("trips") or [])
    n_slots = _FORM_MIN_TRIP_ROWS  # 템플릿 11–22행 = 12칸
    trip_start = 11
    for i in range(n_slots):
        row = trip_start + i
        ws.cell(row=row, column=1, value=i + 1)
        if i < len(trips):
            t = trips[i] or {}
            ws.cell(row=row, column=2, value=t.get("depart_time") or None)
            ws.cell(row=row, column=3, value=t.get("arrive_time") or None)
            ws.cell(row=row, column=4, value=t.get("from") or None)
            ws.cell(row=row, column=5, value=t.get("to") or None)
            ws.cell(row=row, column=6, value=t.get("purpose") or None)
            dist = t.get("distance_km")
            ws.cell(row=row, column=7, value=dist if dist not in ("", None) else None)
            ws.cell(row=row, column=8, value=t.get("duration_display") or None)
            ws.cell(row=row, column=9, value=t.get("memo") or None)
        else:
            for col in range(2, 10):
                ws.cell(row=row, column=col, value=None)

    # 데이터가 12칸을 넘는 경우: 합계 행 앞에 추가 행 삽입
    if len(trips) > n_slots:
        extra = len(trips) - n_slots
        # 합계 행(23) 앞에 삽입
        ws.insert_rows(23, amount=extra)
        for j in range(extra):
            i = n_slots + j
            row = 23 + j
            t = trips[i] or {}
            ws.cell(row=row, column=1, value=i + 1)
            ws.cell(row=row, column=2, value=t.get("depart_time") or None)
            ws.cell(row=row, column=3, value=t.get("arrive_time") or None)
            ws.cell(row=row, column=4, value=t.get("from") or None)
            ws.cell(row=row, column=5, value=t.get("to") or None)
            ws.cell(row=row, column=6, value=t.get("purpose") or None)
            dist = t.get("distance_km")
            ws.cell(row=row, column=7, value=dist if dist not in ("", None) else None)
            ws.cell(row=row, column=8, value=t.get("duration_display") or None)
            ws.cell(row=row, column=9, value=t.get("memo") or None)
            ws.row_dimensions[row].height = 30

    # 합계 행 찾기
    total_row = None
    for r in range(11, ws.max_row + 1):
        if ws.cell(row=r, column=1).value == "합계":
            total_row = r
            break
    if total_row is None:
        total_row = 23 + max(0, len(trips) - n_slots)

    dist_val = log.get("total_distance_km", "")
    if dist_val in ("", None) and trips:
        try:
            dist_val = round(
                sum(float(t.get("distance_km") or 0) for t in trips if t.get("distance_km") not in ("", None)),
                1,
            )
        except (TypeError, ValueError):
            dist_val = None
    ws.cell(row=total_row, column=7, value=dist_val if dist_val not in ("", None) else None)
    ws.cell(row=total_row, column=8, value=meta.get("총 운행시간") or None)

    # 인쇄 설정 유지·보정 (A4 세로 1장)
    try:
        ws.page_setup.paperSize = 9
        ws.page_setup.orientation = "portrait"
        ws.page_setup.fitToWidth = 1
        ws.page_setup.fitToHeight = 1
        if getattr(ws, "sheet_properties", None) is not None:
            if getattr(ws.sheet_properties, "pageSetUpPr", None) is not None:
                ws.sheet_properties.pageSetUpPr.fitToPage = True
        ws.print_area = f"A1:I{ws.max_row}"
    except Exception:
        pass

    # ⚠️ 워터마크/서비스명 없음
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue(), f"{_filename_base(log)}.xlsx"


def _export_excel_built(log: dict) -> tuple[bytes, str]:
    """템플릿 파일 없을 때만 사용하는 코드 생성 폴백 (동일 레이아웃)."""
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
    from openpyxl.utils import get_column_letter

    wb = Workbook()
    ws = wb.active
    ws.title = "운행일지"
    ws.sheet_view.showGridLines = False
    ws.page_setup.paperSize = 9
    ws.page_setup.orientation = "portrait"
    ws.page_setup.fitToPage = True
    ws.page_setup.fitToWidth = 1
    ws.page_setup.fitToHeight = 1
    ws.page_setup.horizontalCentered = True
    ws.page_margins.left = 0.5
    ws.page_margins.right = 0.5
    ws.page_margins.top = 0.45
    ws.page_margins.bottom = 0.45

    thin = Border(
        left=Side(style="thin", color=_FORM_GRID),
        right=Side(style="thin", color=_FORM_GRID),
        top=Side(style="thin", color=_FORM_GRID),
        bottom=Side(style="thin", color=_FORM_GRID),
    )
    head_fill = PatternFill("solid", fgColor=_FORM_HEAD_BG)
    label_fill = PatternFill("solid", fgColor=_FORM_LABEL_BG)
    zebra_fill = PatternFill("solid", fgColor=_FORM_ZEBRA)
    total_fill = PatternFill("solid", fgColor=_FORM_TOTAL_BG)
    white_fill = PatternFill("solid", fgColor="FFFFFF")
    title_font = Font(bold=True, color="FFFFFF", name="맑은 고딕", size=17)
    sub_font = Font(color="CBD5E1", name="맑은 고딕", size=9)
    header_font = Font(bold=True, color="FFFFFF", name="맑은 고딕", size=10)
    label_font = Font(bold=True, color=_FORM_SLATE, name="맑은 고딕", size=9)
    cell_font = Font(name="맑은 고딕", size=10, color=_FORM_NAVY)
    cell_font_sm = Font(name="맑은 고딕", size=9, color=_FORM_NAVY)
    total_font = Font(bold=True, name="맑은 고딕", size=10, color=_FORM_NAVY)
    center = Alignment(horizontal="center", vertical="center", wrap_text=True)
    left = Alignment(horizontal="left", vertical="center", wrap_text=True)

    n_cols = 9
    ws.merge_cells("A1:I1")
    c = ws["A1"]
    c.value = _FORM_TITLE
    c.font = title_font
    c.fill = PatternFill("solid", fgColor=_FORM_NAVY)
    c.alignment = Alignment(horizontal="center", vertical="center")
    ws.row_dimensions[1].height = 34
    ws.merge_cells("A2:I2")
    ws["A2"] = "업무용 차량 운행 기록 · 제출용 · A4"
    ws["A2"].font = sub_font
    ws["A2"].fill = PatternFill("solid", fgColor=_FORM_HEAD_BG)
    ws["A2"].alignment = Alignment(horizontal="center", vertical="center")
    ws.row_dimensions[2].height = 18

    meta = _meta(log)
    meta_rows = _meta_display_rows(log, meta)
    r = 4
    for lab1, val1, lab2, val2 in meta_rows:
        ws.cell(row=r, column=1, value=lab1).font = label_font
        ws.cell(row=r, column=1).fill = label_fill
        ws.cell(row=r, column=1).border = thin
        ws.cell(row=r, column=1).alignment = center
        ws.merge_cells(start_row=r, start_column=2, end_row=r, end_column=4)
        ws.cell(row=r, column=2, value=val1).font = cell_font
        ws.cell(row=r, column=2).alignment = left
        for col in range(2, 5):
            ws.cell(row=r, column=col).border = thin
        ws.cell(row=r, column=5, value=lab2).font = label_font
        ws.cell(row=r, column=5).fill = label_fill
        ws.cell(row=r, column=5).border = thin
        ws.cell(row=r, column=5).alignment = center
        ws.merge_cells(start_row=r, start_column=6, end_row=r, end_column=9)
        ws.cell(row=r, column=6, value=val2).font = cell_font
        for col in range(6, 10):
            ws.cell(row=r, column=col).border = thin
        ws.row_dimensions[r].height = 26
        r += 1
    r += 1
    ws.cell(row=r, column=1, value="운행 요약").font = label_font
    ws.cell(row=r, column=1).fill = label_fill
    ws.cell(row=r, column=1).border = thin
    ws.cell(row=r, column=1).alignment = center
    ws.merge_cells(start_row=r, start_column=2, end_row=r, end_column=9)
    ws.cell(row=r, column=2, value=meta.get("요약") or "").font = cell_font
    for col in range(2, 10):
        ws.cell(row=r, column=col).border = thin
    ws.row_dimensions[r].height = 32
    r += 1
    headers = ["순번", "출발시각", "도착시각", "출발지", "도착지", "운행목적", "거리(km)", "운행시간", "비고"]
    header_row = r
    for i, h in enumerate(headers, start=1):
        cell = ws.cell(row=header_row, column=i, value=h)
        cell.fill = head_fill
        cell.font = header_font
        cell.alignment = center
        cell.border = thin
    ws.row_dimensions[header_row].height = 26
    raw = _trips_dataframe(log).to_dict(orient="records")
    records = _pad_trip_records(raw, _FORM_MIN_TRIP_ROWS)
    for r_idx, rec in enumerate(records, start=1):
        row = header_row + r_idx
        fill = zebra_fill if r_idx % 2 == 0 else white_fill
        for c_idx, h in enumerate(headers, start=1):
            val = rec.get(h, "")
            cell = ws.cell(row=row, column=c_idx, value=val if val != "" else None)
            cell.font = cell_font_sm
            cell.border = thin
            cell.fill = fill
            cell.alignment = center if c_idx in (1, 2, 3, 7, 8) else left
        ws.row_dimensions[row].height = 30
    total_row = header_row + len(records) + 1
    for col in range(1, 10):
        ws.cell(row=total_row, column=col).fill = total_fill
        ws.cell(row=total_row, column=col).border = thin
    ws.cell(row=total_row, column=1, value="합계").font = total_font
    ws.cell(row=total_row, column=1).alignment = center
    ws.cell(row=total_row, column=7, value=log.get("total_distance_km")).font = total_font
    ws.cell(row=total_row, column=7).alignment = center
    ws.cell(row=total_row, column=8, value=meta.get("총 운행시간")).font = total_font
    ws.cell(row=total_row, column=8).alignment = center
    sig_row = total_row + 2
    ws.merge_cells(start_row=sig_row, start_column=5, end_row=sig_row, end_column=6)
    ws.merge_cells(start_row=sig_row, start_column=7, end_row=sig_row, end_column=9)
    ws.merge_cells(start_row=sig_row + 1, start_column=5, end_row=sig_row + 3, end_column=6)
    ws.merge_cells(start_row=sig_row + 1, start_column=7, end_row=sig_row + 3, end_column=9)
    for label, cs, ce in (("작성자", 5, 6), ("확인자", 7, 9)):
        ws.cell(row=sig_row, column=cs, value=label).font = label_font
        ws.cell(row=sig_row, column=cs).fill = label_fill
        ws.cell(row=sig_row, column=cs).alignment = center
        for col in range(cs, ce + 1):
            ws.cell(row=sig_row, column=col).border = thin
            ws.cell(row=sig_row, column=col).fill = label_fill
        ws.cell(row=sig_row + 1, column=cs, value="(서명)").font = Font(
            name="맑은 고딕", size=9, color="94A3B8"
        )
        ws.cell(row=sig_row + 1, column=cs).alignment = center
        for rr in (sig_row + 1, sig_row + 2, sig_row + 3):
            for col in range(cs, ce + 1):
                ws.cell(row=rr, column=col).border = thin
    widths = [5.5, 9.5, 9.5, 17, 17, 12, 8.5, 10.5, 12]
    for i, w in enumerate(widths, start=1):
        ws.column_dimensions[get_column_letter(i)].width = w
    ws.print_area = f"A1:I{sig_row + 3}"
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue(), f"{_filename_base(log)}.xlsx"


# ────────────────────────────────────────────────────────
# PDF (클린 — 워터마크 없음)
# ────────────────────────────────────────────────────────


def _pdf_esc(text: object) -> str:
    """ReportLab Paragraph용 최소 이스케이프."""
    s = str(text or "")
    return (
        s.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace("\n", "<br/>")
    )


def _export_field_pdf(log: dict) -> tuple[bytes, str]:
    """외근·출장 일지 PDF — A4 제출용, 셀 자동 줄바꿈(겹침 방지)."""
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER, TA_LEFT
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import mm
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

    font_name = _register_korean_font()
    page = A4
    margin = 12 * mm
    usable = page[0] - 2 * margin

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=page,
        leftMargin=margin,
        rightMargin=margin,
        topMargin=margin,
        bottomMargin=margin,
        title="외근·출장 일지",
        author=str(log.get("author_name") or log.get("driver_name") or ""),
    )
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "FieldTitle",
        parent=styles["Title"],
        fontName=font_name,
        fontSize=15,
        textColor=colors.white,
        alignment=TA_CENTER,
        leading=19,
        spaceBefore=0,
        spaceAfter=0,
    )
    sub_style = ParagraphStyle(
        "FieldSub",
        parent=styles["Normal"],
        fontName=font_name,
        fontSize=8,
        textColor=colors.HexColor("#CBD5E1"),
        alignment=TA_CENTER,
        leading=11,
    )
    body = ParagraphStyle(
        "FieldBody",
        parent=styles["Normal"],
        fontName=font_name,
        fontSize=9.5,
        leading=13,
        textColor=colors.HexColor("#0F172A"),
        alignment=TA_LEFT,
    )
    label_s = ParagraphStyle(
        "FieldLabel",
        parent=body,
        fontName=font_name,
        fontSize=9,
        textColor=colors.HexColor("#1E293B"),
        alignment=TA_CENTER,
    )
    small = ParagraphStyle(
        "FieldSmall",
        parent=body,
        fontName=font_name,
        fontSize=8.5,
        leading=12,
        textColor=colors.HexColor("#0F172A"),
    )
    head_s = ParagraphStyle(
        "FieldHead",
        parent=body,
        fontName=font_name,
        fontSize=9,
        textColor=colors.white,
        alignment=TA_CENTER,
        leading=12,
    )
    cell_c = ParagraphStyle(
        "FieldCellC",
        parent=small,
        alignment=TA_CENTER,
    )

    author = str(log.get("author_name") or log.get("driver_name") or "—")
    story: list[Any] = []

    # 제목 바
    title_bar = Table(
        [
            [Paragraph("외근·출장 일지", title_style)],
            [Paragraph("방문·업무 결과 기록 · 제출용 · A4", sub_style)],
        ],
        colWidths=[usable],
    )
    title_bar.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor(f"#{_FORM_NAVY}")),
                ("TOPPADDING", (0, 0), (-1, 0), 10),
                ("BOTTOMPADDING", (0, 0), (-1, 0), 2),
                ("TOPPADDING", (0, 1), (-1, 1), 0),
                ("BOTTOMPADDING", (0, 1), (-1, 1), 8),
            ]
        )
    )
    story.append(title_bar)
    story.append(Spacer(1, 3 * mm))

    # 메타 그리드 (라벨|값)
    meta_data = [
        [
            Paragraph("작성일", label_s),
            Paragraph(_pdf_esc(log.get("date") or "—"), body),
            Paragraph("작성자", label_s),
            Paragraph(_pdf_esc(author), body),
        ],
        [
            Paragraph("부서", label_s),
            Paragraph(_pdf_esc(log.get("department") or "—"), body),
            Paragraph("회사명", label_s),
            Paragraph(_pdf_esc(log.get("company_name") or "—"), body),
        ],
    ]
    meta_tbl = Table(
        meta_data,
        colWidths=[usable * 0.14, usable * 0.36, usable * 0.14, usable * 0.36],
        rowHeights=[24, 24],
    )
    meta_tbl.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (0, -1), colors.HexColor(f"#{_FORM_LABEL_BG}")),
                ("BACKGROUND", (2, 0), (2, -1), colors.HexColor(f"#{_FORM_LABEL_BG}")),
                ("BOX", (0, 0), (-1, -1), 0.7, colors.HexColor("#94A3B8")),
                ("INNERGRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#94A3B8")),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (-1, -1), 7),
                ("RIGHTPADDING", (0, 0), (-1, -1), 7),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ]
        )
    )
    story.append(meta_tbl)
    story.append(Spacer(1, 2.5 * mm))

    sum_tbl = Table(
        [
            [
                Paragraph("업무 요약", label_s),
                Paragraph(_pdf_esc(log.get("summary") or "—"), body),
            ]
        ],
        colWidths=[usable * 0.14, usable * 0.86],
        rowHeights=[28],
    )
    sum_tbl.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (0, 0), colors.HexColor(f"#{_FORM_LABEL_BG}")),
                ("BOX", (0, 0), (-1, -1), 0.7, colors.HexColor("#94A3B8")),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (-1, -1), 7),
                ("RIGHTPADDING", (0, 0), (-1, -1), 7),
                ("TOPPADDING", (0, 0), (-1, -1), 6),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
            ]
        )
    )
    story.append(sum_tbl)
    story.append(Spacer(1, 3 * mm))

    # 방문 표 — Paragraph로 자동 줄바꿈 (겹침 방지)
    headers = ["순번", "시각", "방문처", "목적", "결과", "후속"]
    data: list[list[Any]] = [[Paragraph(f"<b>{h}</b>", head_s) for h in headers]]
    visits = _field_visits(log)
    n_rows = max(len(visits), _FORM_MIN_VISIT_ROWS)
    for i in range(n_rows):
        if i < len(visits):
            v = visits[i] or {}
            data.append(
                [
                    Paragraph(str(i + 1), cell_c),
                    Paragraph(_pdf_esc(v.get("time") or ""), cell_c),
                    Paragraph(_pdf_esc(v.get("place") or ""), small),
                    Paragraph(_pdf_esc(v.get("purpose") or ""), small),
                    Paragraph(_pdf_esc(v.get("result") or ""), small),
                    Paragraph(_pdf_esc(v.get("next_action") or ""), small),
                ]
            )
        else:
            data.append(
                [
                    Paragraph(str(i + 1), cell_c),
                    Paragraph(" ", small),
                    Paragraph(" ", small),
                    Paragraph(" ", small),
                    Paragraph(" ", small),
                    Paragraph(" ", small),
                ]
            )

    col_w = [
        usable * 0.07,
        usable * 0.10,
        usable * 0.20,
        usable * 0.18,
        usable * 0.24,
        usable * 0.21,
    ]
    table = Table(data, colWidths=col_w, repeatRows=1)
    style_cmds = [
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(f"#{_FORM_HEAD_BG}")),
        ("FONTNAME", (0, 0), (-1, -1), font_name),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#94A3B8")),
        ("TOPPADDING", (0, 0), (-1, -1), 6),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
        ("LEFTPADDING", (0, 0), (-1, -1), 4),
        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
        ("ALIGN", (0, 0), (1, -1), "CENTER"),
    ]
    for i in range(1, len(data)):
        if i % 2 == 0:
            style_cmds.append(
                ("BACKGROUND", (0, i), (-1, i), colors.HexColor(f"#{_FORM_ZEBRA}"))
            )
    table.setStyle(TableStyle(style_cmds))
    story.append(table)
    story.append(Spacer(1, 6 * mm))

    sig = Table(
        [
            [Paragraph("작성자", label_s), Paragraph("확인자", label_s)],
            [
                Paragraph("<font color='#94A3B8'>( 서 명 )</font>", cell_c),
                Paragraph("<font color='#94A3B8'>( 서 명 )</font>", cell_c),
            ],
        ],
        colWidths=[48 * mm, 48 * mm],
        rowHeights=[14, 28],
    )
    sig.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(f"#{_FORM_LABEL_BG}")),
                ("BOX", (0, 0), (-1, -1), 0.6, colors.HexColor("#94A3B8")),
                ("INNERGRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#94A3B8")),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("ALIGN", (0, 0), (-1, -1), "CENTER"),
            ]
        )
    )
    sig_wrap = Table([[sig]], colWidths=[usable])
    sig_wrap.setStyle(TableStyle([("ALIGN", (0, 0), (-1, -1), "RIGHT")]))
    story.append(sig_wrap)

    doc.build(story)
    return buf.getvalue(), f"{_filename_base(log)}.pdf"


def export_pdf(log: dict) -> tuple[bytes, str]:
    """reportlab PDF. 한글은 시스템 폰트 또는 내장 대체. 회사 제출용 기본 서식."""
    log = _privacy_log(log)
    if _is_field_log(log):
        return _export_field_pdf(log)

    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER, TA_LEFT
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import mm
    from reportlab.platypus import (
        Paragraph,
        SimpleDocTemplate,
        Spacer,
        Table,
        TableStyle,
    )

    font_name = _register_korean_font()
    # A4 세로 (210×297mm) — 인쇄 기본 규격
    page = A4
    margin = 10 * mm

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=page,
        leftMargin=margin,
        rightMargin=margin,
        topMargin=margin,
        bottomMargin=margin,
        title=_FORM_TITLE,
        author=str(log.get("driver_name") or ""),
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "KRTitle",
        parent=styles["Title"],
        fontName=font_name,
        fontSize=16,
        textColor=colors.white,
        alignment=TA_CENTER,
        spaceBefore=0,
        spaceAfter=0,
        leading=20,
    )
    subtitle_style = ParagraphStyle(
        "KRSub",
        parent=styles["Normal"],
        fontName=font_name,
        fontSize=8.5,
        textColor=colors.HexColor("#CBD5E1"),
        alignment=TA_CENTER,
        leading=11,
    )
    body = ParagraphStyle(
        "KRBody",
        parent=styles["Normal"],
        fontName=font_name,
        fontSize=9.5,
        leading=13,
        textColor=colors.HexColor("#0F172A"),
    )
    label_s = ParagraphStyle(
        "KRLabel",
        parent=body,
        fontName=font_name,
        fontSize=9,
        textColor=colors.HexColor("#1E293B"),
        alignment=TA_CENTER,
    )
    small = ParagraphStyle(
        "KRSmall",
        parent=body,
        fontSize=9,
        leading=12,
        textColor=colors.HexColor("#0F172A"),
    )
    head_s = ParagraphStyle(
        "KRHead",
        parent=body,
        fontSize=9,
        textColor=colors.white,
        alignment=TA_CENTER,
        leading=12,
    )

    meta = _meta(log)
    story: list[Any] = []
    usable_w = page[0] - 2 * margin

    # 제목 바
    title_bar = Table(
        [
            [Paragraph(_FORM_TITLE, title_style)],
            [Paragraph("업무용 차량 운행 기록 · 제출용 · A4", subtitle_style)],
        ],
        colWidths=[usable_w],
    )
    title_bar.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor(f"#{_FORM_NAVY}")),
                ("TOPPADDING", (0, 0), (-1, 0), 10),
                ("BOTTOMPADDING", (0, 0), (-1, 0), 2),
                ("TOPPADDING", (0, 1), (-1, 1), 0),
                ("BOTTOMPADDING", (0, 1), (-1, 1), 8),
                ("ALIGN", (0, 0), (-1, -1), "CENTER"),
            ]
        )
    )
    story.append(title_bar)
    story.append(Spacer(1, 3 * mm))

    # 메타 그리드 (라벨|값|라벨|값) — 라벨 칸 넓게
    meta_rows = _meta_display_rows(log, meta)
    meta_table_data = []
    for lab1, val1, lab2, val2 in meta_rows:
        meta_table_data.append(
            [
                Paragraph(lab1, label_s),
                Paragraph(str(val1), body),
                Paragraph(lab2, label_s),
                Paragraph(str(val2), body),
            ]
        )
    usable = page[0] - 2 * margin  # A4 폭 − 좌우 여백
    meta_table = Table(
        meta_table_data,
        colWidths=[usable * 0.18, usable * 0.32, usable * 0.18, usable * 0.32],
        rowHeights=[22] * len(meta_table_data),
    )
    meta_table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (0, -1), colors.HexColor(f"#{_FORM_LABEL_BG}")),
                ("BACKGROUND", (2, 0), (2, -1), colors.HexColor(f"#{_FORM_LABEL_BG}")),
                ("BOX", (0, 0), (-1, -1), 0.7, colors.HexColor("#94A3B8")),
                ("INNERGRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#94A3B8")),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (-1, -1), 8),
                ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                ("TOPPADDING", (0, 0), (-1, -1), 6),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
            ]
        )
    )
    story.append(meta_table)
    story.append(Spacer(1, 2.5 * mm))

    # 요약
    summary_tbl = Table(
        [
            [
                Paragraph("운행 요약", label_s),
                Paragraph(str(meta.get("요약") or "—"), body),
            ]
        ],
        colWidths=[usable * 0.18, usable * 0.82],
        rowHeights=[28],
    )
    summary_tbl.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (0, 0), colors.HexColor(f"#{_FORM_LABEL_BG}")),
                ("BOX", (0, 0), (-1, -1), 0.7, colors.HexColor("#94A3B8")),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (-1, -1), 8),
                ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                ("TOPPADDING", (0, 0), (-1, -1), 6),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
            ]
        )
    )
    story.append(summary_tbl)
    story.append(Spacer(1, 2 * mm))

    # 본문 테이블 — 최소 행 수만큼 빈칸 채워 A4 표를 꽉 채움
    headers = ["순번", "출발시각", "도착시각", "출발지", "도착지", "목적", "km", "운행시간", "비고"]
    data = [[Paragraph(f"<b>{h}</b>", head_s) for h in headers]]
    trips_src = list(log.get("trips") or [])
    n_rows = max(len(trips_src), _FORM_MIN_TRIP_ROWS)
    for i in range(n_rows):
        if i < len(trips_src):
            t = trips_src[i] or {}
            data.append(
                [
                    Paragraph(str(i + 1), small),
                    Paragraph(str(t.get("depart_time", "")), small),
                    Paragraph(str(t.get("arrive_time", "")), small),
                    Paragraph(str(t.get("from", "")), small),
                    Paragraph(str(t.get("to", "")), small),
                    Paragraph(str(t.get("purpose", "")), small),
                    Paragraph(str(t.get("distance_km", "")), small),
                    Paragraph(str(t.get("duration_display", "")), small),
                    Paragraph(str(t.get("memo", "")), small),
                ]
            )
        else:
            data.append(
                [
                    Paragraph(str(i + 1), small),
                    Paragraph(" ", small),
                    Paragraph(" ", small),
                    Paragraph(" ", small),
                    Paragraph(" ", small),
                    Paragraph(" ", small),
                    Paragraph(" ", small),
                    Paragraph(" ", small),
                    Paragraph(" ", small),
                ]
            )
    # 합계 행
    data.append(
        [
            Paragraph("<b>합계</b>", small),
            Paragraph("", small),
            Paragraph("", small),
            Paragraph("", small),
            Paragraph("", small),
            Paragraph("", small),
            Paragraph(f"<b>{meta.get('총 거리(km)') or ''}</b>", small),
            Paragraph(f"<b>{meta.get('총 운행시간') or ''}</b>", small),
            Paragraph("", small),
        ]
    )

    # A4 세로 본문폭에 맞춘 열 비율 (합 1.0)
    col_w = [
        usable * 0.05,
        usable * 0.08,
        usable * 0.08,
        usable * 0.17,
        usable * 0.17,
        usable * 0.12,
        usable * 0.07,
        usable * 0.10,
        usable * 0.16,
    ]
    # 남는 세로 공간을 구간 행 높이에 분배 → 빈칸이 있어도 A4 1장을 꽉 채움
    usable_h = page[1] - 2 * margin
    header_block = 78 * mm
    footer_block = 34 * mm
    trip_area = max(usable_h - header_block - footer_block, 90 * mm)
    row_h = float(trip_area) / max(n_rows, 1)
    row_h = min(max(row_h, 14), 24)
    row_heights = [18] + [row_h] * n_rows + [18]
    table = Table(data, colWidths=col_w, rowHeights=row_heights, repeatRows=1)
    style_cmds = [
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(f"#{_FORM_HEAD_BG}")),
        ("BACKGROUND", (0, -1), (-1, -1), colors.HexColor(f"#{_FORM_TOTAL_BG}")),
        ("FONTNAME", (0, 0), (-1, -1), font_name),
        ("ALIGN", (0, 0), (-1, 0), "CENTER"),
        ("ALIGN", (0, 1), (0, -1), "CENTER"),
        ("ALIGN", (1, 1), (2, -1), "CENTER"),
        ("ALIGN", (6, 1), (7, -1), "CENTER"),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#94A3B8")),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ("LEFTPADDING", (0, 0), (-1, -1), 4),
        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
    ]
    for i in range(1, len(data) - 1):
        if i % 2 == 0:
            style_cmds.append(
                ("BACKGROUND", (0, i), (-1, i), colors.HexColor(f"#{_FORM_ZEBRA}"))
            )
    table.setStyle(TableStyle(style_cmds))
    story.append(table)
    story.append(Spacer(1, 3 * mm))

    # 서명 박스 (하단)
    sig = Table(
        [
            [
                Paragraph("작성자", label_s),
                Paragraph("확인자", label_s),
            ],
            [
                Paragraph("<font color='#94A3B8'>(서명)</font>", small),
                Paragraph("<font color='#94A3B8'>(서명)</font>", small),
            ],
        ],
        colWidths=[48 * mm, 48 * mm],
        rowHeights=[11, 18],
    )
    sig.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(f"#{_FORM_LABEL_BG}")),
                ("BOX", (0, 0), (-1, -1), 0.6, colors.HexColor(f"#{_FORM_GRID}")),
                ("INNERGRID", (0, 0), (-1, -1), 0.4, colors.HexColor(f"#{_FORM_GRID}")),
                ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("TOPPADDING", (0, 0), (-1, 0), 3),
                ("BOTTOMPADDING", (0, 0), (-1, 0), 3),
                ("TOPPADDING", (0, 1), (-1, 1), 8),
                ("BOTTOMPADDING", (0, 1), (-1, 1), 8),
            ]
        )
    )
    # 오른쪽 정렬용 래퍼
    sig_wrap = Table([[sig]], colWidths=[usable])
    sig_wrap.setStyle(TableStyle([("ALIGN", (0, 0), (-1, -1), "RIGHT")]))
    story.append(sig_wrap)
    # ⚠️ 워터마크 / "Powered by" / 서비스 로고 텍스트 없음

    doc.build(story)
    return buf.getvalue(), f"{_filename_base(log)}.pdf"


_KR_FONT_REGISTERED: str | None = None


def _register_korean_font() -> str:
    """Windows/Mac/Linux 한글 폰트 등록. 실패 시 Helvetica(한글 깨짐)."""
    global _KR_FONT_REGISTERED
    if _KR_FONT_REGISTERED:
        return _KR_FONT_REGISTERED

    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    import os
    from pathlib import Path

    root = Path(__file__).resolve().parents[1]
    candidates = [
        # 앱 번들 (Docker/로컬 공통)
        str(root / "assets" / "fonts" / "NanumGothic.ttf"),
        str(root / "web" / "assets" / "fonts" / "NanumGothic.ttf"),
        # Windows
        r"C:\Windows\Fonts\malgun.ttf",
        r"C:\Windows\Fonts\malgunbd.ttf",
        r"C:\Windows\Fonts\NanumGothic.ttf",
        r"C:\Windows\Fonts\NanumBarunGothic.ttf",
        # macOS
        "/System/Library/Fonts/AppleSDGothicNeo.ttc",
        "/Library/Fonts/AppleGothic.ttf",
        # Linux (apt fonts-nanum / fonts-noto-cjk)
        "/usr/share/fonts/truetype/nanum/NanumGothic.ttf",
        "/usr/share/fonts/truetype/nanum/NanumBarunGothic.ttf",
        "/usr/share/fonts/truetype/nanum/NanumGothicCoding.ttf",
        "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
        "/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc",
        "/usr/share/fonts/truetype/noto/NotoSansKR-Regular.otf",
        "/usr/share/fonts/opentype/noto/NotoSansKR-Regular.otf",
    ]
    for path in candidates:
        if not os.path.exists(path):
            continue
        try:
            name = "KRFont"
            # 이미 등록된 경우 재사용
            if name in pdfmetrics.getRegisteredFontNames():
                _KR_FONT_REGISTERED = name
                return name
            if path.lower().endswith(".ttc"):
                pdfmetrics.registerFont(TTFont(name, path, subfontIndex=0))
            else:
                pdfmetrics.registerFont(TTFont(name, path))
            _KR_FONT_REGISTERED = name
            return name
        except Exception:
            continue
    _KR_FONT_REGISTERED = "Helvetica"
    return "Helvetica"


# ────────────────────────────────────────────────────────
# DOCX (한컴 오피스 호환 — 클린)
# ────────────────────────────────────────────────────────


def _export_field_docx(log: dict) -> tuple[bytes, str]:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor

    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)

    def set_run_font(run, size=10, bold=False, color=None, font="맑은 고딕"):
        run.bold = bold
        run.font.size = Pt(size)
        run.font.name = font
        r = run._element
        r.rPr.rFonts.set(qn("w:eastAsia"), font)
        if color:
            run.font.color.rgb = RGBColor(*color)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("외근·출장 일지")
    set_run_font(tr, size=18, bold=True, color=(11, 31, 58))

    author = str(log.get("author_name") or log.get("driver_name") or "")
    info = doc.add_paragraph()
    ir = info.add_run(
        f"작성일: {log.get('date') or ''}    작성자: {author}    "
        f"부서: {log.get('department') or ''}    회사명: {log.get('company_name') or ''}"
    )
    set_run_font(ir, size=10)

    sum_p = doc.add_paragraph()
    sr = sum_p.add_run(f"업무 요약: {log.get('summary') or ''}")
    set_run_font(sr, size=10, bold=True)

    headers = ["순번", "시각", "방문처", "목적", "결과", "후속 조치", "비고"]
    visits = _field_visits(log)
    table = doc.add_table(rows=1 + max(len(visits), 1), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        set_run_font(run, size=9, bold=True, color=(11, 31, 58))

    rows = visits or [{"time": "", "place": "", "purpose": "", "result": "", "next_action": "", "memo": ""}]
    for r_idx, v in enumerate(rows):
        vals = [
            str(r_idx + 1),
            str(v.get("time") or ""),
            str(v.get("place") or ""),
            str(v.get("purpose") or ""),
            str(v.get("result") or ""),
            str(v.get("next_action") or ""),
            str(v.get("memo") or ""),
        ]
        for c_idx, val in enumerate(vals):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(val)
            set_run_font(run, size=9)

    sig = doc.add_paragraph()
    set_run_font(sig.add_run("\n작성자: _______________    확인자: _______________"), size=10)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue(), f"{_filename_base(log)}.docx"


def export_docx(log: dict) -> tuple[bytes, str]:
    """
    python-docx 기반 .docx
    한글과 컴퓨터(한컴) 오피스에서 정상 개방되는 표준 OOXML.
    회사 제출용 기본 서식 · 워터마크/머리글 광고 없음.
    """
    log = _privacy_log(log)
    if _is_field_log(log):
        return _export_field_docx(log)

    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor

    doc = Document()

    for section in doc.sections:
        try:
            from docx.enum.section import WD_ORIENT

            section.orientation = WD_ORIENT.PORTRAIT
        except Exception:
            pass
        # A4 세로 210×297mm
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(1.0)
        section.bottom_margin = Cm(1.0)
        section.left_margin = Cm(1.2)
        section.right_margin = Cm(1.2)

    def set_run_font(run, size=10, bold=False, color=None, font="맑은 고딕"):
        run.bold = bold
        run.font.size = Pt(size)
        run.font.name = font
        r = run._element
        r.rPr.rFonts.set(qn("w:eastAsia"), font)
        if color:
            run.font.color.rgb = RGBColor(*color)

    def set_cell_text(cell, text, *, size=9, bold=False, color=None, align="left"):
        cell.text = ""
        p = cell.paragraphs[0]
        if align == "center":
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif align == "right":
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(str(text if text is not None else ""))
        set_run_font(run, size=size, bold=bold, color=color)

    # 제목 바
    title_tbl = doc.add_table(rows=2, cols=1)
    title_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_cell_text(
        title_tbl.rows[0].cells[0],
        _FORM_TITLE,
        size=16,
        bold=True,
        color=(255, 255, 255),
        align="center",
    )
    _set_cell_shading(title_tbl.rows[0].cells[0], _FORM_NAVY)
    set_cell_text(
        title_tbl.rows[1].cells[0],
        "업무용 차량 운행 기록 · 제출용",
        size=8,
        color=(203, 213, 225),
        align="center",
    )
    _set_cell_shading(title_tbl.rows[1].cells[0], _FORM_HEAD_BG)

    doc.add_paragraph()

    meta = _meta(log)
    meta_rows = _meta_display_rows(log, meta)
    meta_tbl = doc.add_table(rows=len(meta_rows), cols=4)
    meta_tbl.style = "Table Grid"
    for i, (lab1, val1, lab2, val2) in enumerate(meta_rows):
        set_cell_text(meta_tbl.rows[i].cells[0], lab1, size=8, bold=True, color=(51, 65, 85), align="center")
        set_cell_text(meta_tbl.rows[i].cells[1], val1, size=9)
        set_cell_text(meta_tbl.rows[i].cells[2], lab2, size=8, bold=True, color=(51, 65, 85), align="center")
        set_cell_text(meta_tbl.rows[i].cells[3], val2, size=9)
        _set_cell_shading(meta_tbl.rows[i].cells[0], _FORM_LABEL_BG)
        _set_cell_shading(meta_tbl.rows[i].cells[2], _FORM_LABEL_BG)

    doc.add_paragraph()

    # 요약
    sum_tbl = doc.add_table(rows=1, cols=2)
    sum_tbl.style = "Table Grid"
    set_cell_text(sum_tbl.rows[0].cells[0], "운행 요약", size=8, bold=True, color=(51, 65, 85), align="center")
    set_cell_text(sum_tbl.rows[0].cells[1], meta.get("요약") or "", size=9)
    _set_cell_shading(sum_tbl.rows[0].cells[0], _FORM_LABEL_BG)

    doc.add_paragraph()

    headers = [
        "순번",
        "출발시각",
        "도착시각",
        "출발지",
        "도착지",
        "운행목적",
        "거리(km)",
        "운행시간",
        "비고",
    ]
    trips = list(log.get("trips") or [])
    n_rows = max(len(trips), _FORM_MIN_TRIP_ROWS)
    table = doc.add_table(rows=1 + n_rows + 1, cols=len(headers))
    table.style = "Table Grid"

    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, size=8, bold=True, color=(255, 255, 255), align="center")
        _set_cell_shading(table.rows[0].cells[i], _FORM_HEAD_BG)

    for r_idx in range(n_rows):
        t = trips[r_idx] if r_idx < len(trips) else {}
        vals = [
            str(r_idx + 1),
            str(t.get("depart_time", "")),
            str(t.get("arrive_time", "")),
            str(t.get("from", "")),
            str(t.get("to", "")),
            str(t.get("purpose", "")),
            str(t.get("distance_km", "") if t else ""),
            str(t.get("duration_display", "")),
            str(t.get("memo", "")),
        ]
        for c_idx, val in enumerate(vals):
            align = "center" if c_idx in (0, 1, 2, 6, 7) else "left"
            set_cell_text(table.rows[r_idx + 1].cells[c_idx], val, size=8, align=align)
            if (r_idx + 1) % 2 == 0:
                _set_cell_shading(table.rows[r_idx + 1].cells[c_idx], _FORM_ZEBRA)

    # 합계
    total_row = table.rows[1 + n_rows]
    set_cell_text(total_row.cells[0], "합계", size=9, bold=True, align="center")
    set_cell_text(total_row.cells[6], str(meta.get("총 거리(km)") or ""), size=9, bold=True, align="center")
    set_cell_text(total_row.cells[7], str(meta.get("총 운행시간") or ""), size=9, bold=True, align="center")
    for c in total_row.cells:
        _set_cell_shading(c, _FORM_TOTAL_BG)

    doc.add_paragraph()

    # 서명
    sig_tbl = doc.add_table(rows=2, cols=2)
    sig_tbl.style = "Table Grid"
    set_cell_text(sig_tbl.rows[0].cells[0], "작성자", size=8, bold=True, align="center")
    set_cell_text(sig_tbl.rows[0].cells[1], "확인자", size=8, bold=True, align="center")
    set_cell_text(sig_tbl.rows[1].cells[0], "(서명)", size=9, color=(148, 163, 184), align="center")
    set_cell_text(sig_tbl.rows[1].cells[1], "(서명)", size=9, color=(148, 163, 184), align="center")
    _set_cell_shading(sig_tbl.rows[0].cells[0], _FORM_LABEL_BG)
    _set_cell_shading(sig_tbl.rows[0].cells[1], _FORM_LABEL_BG)

    # ⚠️ 워터마크 / 머리글 / 바닥글 서비스 문구 없음 (한컴 호환 클린 문서)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue(), f"{_filename_base(log)}.docx"


def _set_cell_shading(cell, hex_color: str) -> None:
    """셀 배경색 설정 (OOXML)."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


# ────────────────────────────────────────────────────────
# 메모리 전용 빌더 (서버 디스크 저장 없음)
# ────────────────────────────────────────────────────────


def build_excel_bytes(log: dict) -> tuple[bytes, str]:
    """
    Excel (.xlsx) — 메모리에서 생성.
    openpyxl 서식 + pandas DataFrame 구간 데이터 기반.
    디스크에 파일을 쓰지 않음.
    """
    return export_excel(log)


def build_excel_bytes_pandas(log: dict) -> tuple[bytes, str]:
    """
    pandas 전용 간단 Excel (참고용/폴백).
    BytesIO 만 사용 — 서버 저장 없음.
    """
    import io

    df = _trips_dataframe(log)
    meta = _meta(log)
    buf = io.BytesIO()
    with pd.ExcelWriter(buf, engine="openpyxl") as writer:
        # 요약 시트
        meta_df = pd.DataFrame([{"항목": k, "내용": v} for k, v in meta.items()])
        meta_df.to_excel(writer, sheet_name="요약", index=False)
        # 구간 시트
        if df.empty:
            pd.DataFrame(columns=["순번", "출발시각", "도착시각", "출발지", "도착지"]).to_excel(
                writer, sheet_name="운행구간", index=False
            )
        else:
            df.to_excel(writer, sheet_name="운행구간", index=False)
    return buf.getvalue(), f"{_filename_base(log)}.xlsx"


def export_summary_excel(summary: dict[str, Any]) -> tuple[bytes, str]:
    """주간/월간 업무 요약 → Excel (제출·메일 첨부용)."""
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Font, PatternFill

    s = summary or {}
    wb = Workbook()
    ws = wb.active
    ws.title = "업무요약"

    header_fill = PatternFill("solid", fgColor="0F172A")
    header_font = Font(color="FFFFFF", bold=True, size=11)
    title_font = Font(bold=True, size=14, color="0F172A")
    accent = Font(bold=True, color="0D9488")

    ws["A1"] = "로드로그 업무 요약"
    ws["A1"].font = title_font
    ws.merge_cells("A1:D1")
    ws["A2"] = f"{s.get('period_label') or ''} · {s.get('date_from') or ''} ~ {s.get('date_to') or ''}"
    ws["A2"].font = Font(color="64748B", size=10)
    ws.merge_cells("A2:D2")

    rows = [
        ("총 일지(건)", s.get("total_logs") or 0),
        ("운행일지(건)", s.get("driving_count") or 0),
        ("외근·출장(건)", s.get("field_count") or 0),
        ("총 운행거리(km)", s.get("total_distance_km") or 0),
    ]
    ws["A4"] = "지표"
    ws["B4"] = "값"
    for col in ("A", "B"):
        cell = ws[f"{col}4"]
        cell.fill = header_fill
        cell.font = header_font
    for i, (k, v) in enumerate(rows, start=5):
        ws[f"A{i}"] = k
        ws[f"B{i}"] = v
        ws[f"B{i}"].font = accent

    ws["A10"] = "주요 방문지 Top 3"
    ws["A10"].font = Font(bold=True, size=12)
    ws["A11"] = "순위"
    ws["B11"] = "방문지"
    ws["C11"] = "횟수"
    for col in ("A", "B", "C"):
        cell = ws[f"{col}11"]
        cell.fill = header_fill
        cell.font = header_font
    tops = s.get("top_places") or []
    if not tops:
        ws["A12"] = "—"
        ws["B12"] = "(기록 없음)"
        ws["C12"] = 0
    else:
        for i, tp in enumerate(tops[:3], start=12):
            ws[f"A{i}"] = i - 11
            ws[f"B{i}"] = tp.get("place") or ""
            ws[f"C{i}"] = tp.get("count") or 0

    ws2 = wb.create_sheet("일지목록")
    headers = ["날짜", "유형", "거리(km)", "제목", "요약"]
    for c, h in enumerate(headers, start=1):
        cell = ws2.cell(1, c, h)
        cell.fill = header_fill
        cell.font = header_font
    for r, row in enumerate(s.get("logs") or [], start=2):
        kind = "외근" if row.get("report_type") == "field" else "운행"
        ws2.cell(r, 1, row.get("date") or "")
        ws2.cell(r, 2, kind)
        ws2.cell(r, 3, row.get("distance_km") or 0)
        ws2.cell(r, 4, row.get("title") or "")
        ws2.cell(r, 5, row.get("summary") or "")

    ws3 = wb.create_sheet("제출용텍스트")
    ws3["A1"] = s.get("report_text") or ""
    ws3["A1"].alignment = Alignment(wrap_text=True, vertical="top")
    ws3.column_dimensions["A"].width = 80
    ws3.row_dimensions[1].height = 220

    for col, w in (("A", 22), ("B", 36), ("C", 12), ("D", 24)):
        ws.column_dimensions[col].width = w
    for col, w in (("A", 12), ("B", 10), ("C", 12), ("D", 22), ("E", 48)):
        ws2.column_dimensions[col].width = w

    buf = io.BytesIO()
    wb.save(buf)
    period = s.get("period") or "summary"
    d0 = str(s.get("date_from") or "").replace("-", "")
    d1 = str(s.get("date_to") or "").replace("-", "")
    name = f"RoadLog_업무요약_{period}_{d0}_{d1}.xlsx"
    return buf.getvalue(), name


def export_summary_pdf(summary: dict[str, Any]) -> tuple[bytes, str]:
    """주간/월간 업무 요약 → PDF (결재판·메일 첨부용, 워터마크 없음)."""
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER, TA_LEFT
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import mm
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

    s = summary or {}
    font_name = _register_korean_font()
    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        leftMargin=16 * mm,
        rightMargin=16 * mm,
        topMargin=16 * mm,
        bottomMargin=16 * mm,
        title="로드로그 업무 요약",
        author="RoadLog",
    )
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "SumTitle",
        parent=styles["Title"],
        fontName=font_name,
        fontSize=18,
        textColor=colors.HexColor("#0B1F3A"),
        alignment=TA_CENTER,
        spaceAfter=8,
    )
    body = ParagraphStyle(
        "SumBody",
        parent=styles["Normal"],
        fontName=font_name,
        fontSize=10,
        leading=15,
        textColor=colors.HexColor("#0F172A"),
        alignment=TA_LEFT,
    )
    small = ParagraphStyle(
        "SumSmall",
        parent=body,
        fontSize=9,
        textColor=colors.HexColor("#475569"),
    )

    story: list[Any] = []
    story.append(Paragraph("로드로그 업무 요약", title_style))
    story.append(
        Paragraph(
            f"{s.get('period_label') or ''} &nbsp;|&nbsp; "
            f"{s.get('date_from') or ''} ~ {s.get('date_to') or ''}",
            small,
        )
    )
    story.append(Spacer(1, 5 * mm))

    metrics = [
        [
            Paragraph(f"<b>총 일지</b><br/>{s.get('total_logs') or 0}건", body),
            Paragraph(f"<b>운행</b><br/>{s.get('driving_count') or 0}건", body),
            Paragraph(f"<b>외근</b><br/>{s.get('field_count') or 0}건", body),
            Paragraph(
                f"<b>총 거리</b><br/>{s.get('total_distance_km') or 0} km", body
            ),
        ]
    ]
    mt = Table(metrics, colWidths=[42 * mm, 42 * mm, 42 * mm, 42 * mm])
    mt.setStyle(
        TableStyle(
            [
                ("FONTNAME", (0, 0), (-1, -1), font_name),
                ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#F0FDFA")),
                ("BOX", (0, 0), (-1, -1), 0.8, colors.HexColor("#0D9488")),
                ("INNERGRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#99F6E4")),
                ("TOPPADDING", (0, 0), (-1, -1), 8),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ]
        )
    )
    story.append(mt)
    story.append(Spacer(1, 6 * mm))

    story.append(Paragraph("<b>주요 방문지 Top 3</b>", body))
    story.append(Spacer(1, 2 * mm))
    tops = s.get("top_places") or []
    top_data = [["순위", "방문지", "횟수"]]
    if not tops:
        top_data.append(["—", "(기록 없음)", "0"])
    else:
        for i, tp in enumerate(tops[:3], start=1):
            top_data.append([str(i), str(tp.get("place") or ""), str(tp.get("count") or 0)])
    tt = Table(top_data, colWidths=[20 * mm, 120 * mm, 28 * mm])
    tt.setStyle(
        TableStyle(
            [
                ("FONTNAME", (0, 0), (-1, -1), font_name),
                ("FONTSIZE", (0, 0), (-1, -1), 9),
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#0B1F3A")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#CBD5E1")),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                ("LEFTPADDING", (0, 0), (-1, -1), 5),
            ]
        )
    )
    story.append(tt)
    story.append(Spacer(1, 6 * mm))

    story.append(Paragraph("<b>기간 내 일지</b>", body))
    story.append(Spacer(1, 2 * mm))
    log_data = [["날짜", "유형", "거리", "요약"]]
    for row in (s.get("logs") or [])[:40]:
        kind = "외근" if row.get("report_type") == "field" else "운행"
        km = row.get("distance_km") or ""
        if km != "" and km is not None:
            try:
                km = f"{float(km):g} km"
            except (TypeError, ValueError):
                km = str(km)
        summary = str(row.get("summary") or row.get("title") or "")[:60]
        log_data.append(
            [
                str(row.get("date") or ""),
                kind,
                str(km),
                summary,
            ]
        )
    if len(log_data) == 1:
        log_data.append(["—", "—", "—", "(일지 없음)"])
    lt = Table(log_data, colWidths=[24 * mm, 16 * mm, 22 * mm, 106 * mm])
    lt.setStyle(
        TableStyle(
            [
                ("FONTNAME", (0, 0), (-1, -1), font_name),
                ("FONTSIZE", (0, 0), (-1, -1), 8),
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#0F172A")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("GRID", (0, 0), (-1, -1), 0.3, colors.HexColor("#CBD5E1")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                ("LEFTPADDING", (0, 0), (-1, -1), 4),
            ]
        )
    )
    story.append(lt)
    story.append(Spacer(1, 8 * mm))
    story.append(
        Paragraph(
            "본 문서는 RoadLog에 저장된 일지를 집계한 제출용 요약입니다. "
            "상세 일지는 각 일자 문서를 함께 첨부해 주세요.",
            small,
        )
    )
    story.append(Paragraph("https://roadlog.co.kr", small))

    doc.build(story)
    period = s.get("period") or "summary"
    d0 = str(s.get("date_from") or "").replace("-", "")
    d1 = str(s.get("date_to") or "").replace("-", "")
    name = f"RoadLog_업무요약_{period}_{d0}_{d1}.pdf"
    return buf.getvalue(), name


def build_pdf_bytes(log: dict) -> tuple[bytes, str]:
    """PDF — reportlab, 한글 폰트(맑은고딕) 자동, 메모리 전용."""
    return export_pdf(log)


def build_docx_bytes(log: dict) -> tuple[bytes, str]:
    """DOCX — python-docx 회사 양식형, 메모리 전용."""
    return export_docx(log)


def build_all_download_files(log: dict) -> dict[str, Any]:
    """
    세 형식을 한 번에 메모리에서 생성.
    반환: {
      "excel": (bytes, filename) | None,
      "pdf": ...,
      "docx": ...,
      "errors": {"excel": str, ...}
    }
    """
    out: dict[str, Any] = {
        "excel": None,
        "pdf": None,
        "docx": None,
        "errors": {},
    }
    try:
        out["excel"] = build_excel_bytes(log)
    except Exception as e:
        # 폴백: pandas 간단 엑셀
        try:
            out["excel"] = build_excel_bytes_pandas(log)
        except Exception as e2:
            out["errors"]["excel"] = f"{e} / fallback: {e2}"
    try:
        out["pdf"] = build_pdf_bytes(log)
    except Exception as e:
        out["errors"]["pdf"] = str(e)
    try:
        out["docx"] = build_docx_bytes(log)
    except Exception as e:
        out["errors"]["docx"] = str(e)
    return out


# ────────────────────────────────────────────────────────
# Streamlit 다운로드 버튼 3종 (반응형 · 메모리 · spinner)
# ────────────────────────────────────────────────────────


def _log_fingerprint(log: dict) -> str:
    """세션 캐시 키용 — 동일 log면 재생성 생략."""
    import hashlib
    import json

    raw = json.dumps(log, ensure_ascii=False, sort_keys=True, default=str)
    return hashlib.md5(raw.encode("utf-8")).hexdigest()


def render_download_buttons(log: dict) -> None:
    """
    Excel / PDF / DOCX 다운로드 UI.

    - 서버 디스크 저장 없음 (bytes 메모리)
    - PC: 가로 3열 / 모바일: CSS로 세로 스택
    - st.spinner 로 생성 중 표시 후 버튼 활성화
    """
    import streamlit as st

    st.markdown(
        """
<div class="rl-download-zone">
  <p class="rl-dz-title">문서 다운로드</p>
  <p class="rl-dz-cap">서버에 저장하지 않습니다 · 브라우저로 바로 다운로드 · Excel · PDF · DOCX</p>
</div>
        """,
        unsafe_allow_html=True,
    )

    if not log or not isinstance(log, dict):
        st.info("다운로드할 운행일지 데이터가 없습니다. 먼저 일지를 생성해 주세요.")
        return

    fp = _log_fingerprint(log)
    cache_key = "download_bundle"
    cache_fp_key = "download_bundle_fp"

    # 동일 결과면 재생성 없이 캐시 사용
    need_build = (
        st.session_state.get(cache_fp_key) != fp
        or not st.session_state.get(cache_key)
    )

    if need_build:
        with st.spinner("Excel · PDF · DOCX 파일을 준비하는 중… (서버에 저장하지 않습니다)"):
            bundle = build_all_download_files(log)
        st.session_state[cache_key] = bundle
        st.session_state[cache_fp_key] = fp
    else:
        bundle = st.session_state[cache_key]

    errors = bundle.get("errors") or {}
    excel = bundle.get("excel")
    pdf = bundle.get("pdf")
    docx = bundle.get("docx")

    ready = sum(1 for x in (excel, pdf, docx) if x)
    if ready == 3:
        st.success("다운로드 준비가 완료되었습니다. 원하는 형식을 눌러 받으세요.")
    elif ready > 0:
        st.warning(f"일부 형식만 준비되었습니다. ({ready}/3)")
    else:
        st.error("파일 생성에 실패했습니다.")

    # PC 가로 3열 / 모바일 CSS 스택 (styles.py @media)
    c1, c2, c3 = st.columns(3)

    with c1:
        if excel:
            data, name = excel
            st.download_button(
                label="📊 Excel 다운로드",
                data=data,
                file_name=name,
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                use_container_width=True,
                key=f"dl_excel_{fp[:8]}",
                type="primary",
            )
        else:
            st.button("📊 Excel 실패", disabled=True, use_container_width=True, key="dl_excel_dis")
            if errors.get("excel"):
                st.caption(errors["excel"][:120])

    with c2:
        if pdf:
            data, name = pdf
            st.download_button(
                label="📕 PDF 다운로드",
                data=data,
                file_name=name,
                mime="application/pdf",
                use_container_width=True,
                key=f"dl_pdf_{fp[:8]}",
                type="primary",
            )
        else:
            st.button("📕 PDF 실패", disabled=True, use_container_width=True, key="dl_pdf_dis")
            if errors.get("pdf"):
                st.caption(errors["pdf"][:120])

    with c3:
        if docx:
            data, name = docx
            st.download_button(
                label="📝 DOCX 다운로드",
                data=data,
                file_name=name,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
                key=f"dl_docx_{fp[:8]}",
                type="primary",
            )
        else:
            st.button("📝 DOCX 실패", disabled=True, use_container_width=True, key="dl_docx_dis")
            if errors.get("docx"):
                st.caption(errors["docx"][:120])

    st.caption(
        "파일은 메모리에서만 생성되며 서버 디스크에 저장되지 않습니다. "
        "모바일에서는 버튼이 세로로 배치됩니다."
    )


def clear_download_cache() -> None:
    """일지 재생성 시 호출 — 세션 다운로드 캐시 초기화."""
    try:
        import streamlit as st

        for k in ("download_bundle", "download_bundle_fp"):
            if k in st.session_state:
                del st.session_state[k]
    except Exception:
        pass

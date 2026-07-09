"""
사용자별 운행일지 서식·말투 학습
- 기존 작성 일지 3~5장 업로드
- 텍스트 추출 → 스타일 프로필 생성
- 생성 시 프롬프트에 주입해 '그 사람이 쓴 것처럼' 작성
"""

from __future__ import annotations

import json
import re
import secrets
import uuid
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from modules.config import DATA_DIR, OPENAI_API_KEY, OPENAI_MODEL

STYLES_DIR = DATA_DIR / "styles"
STYLES_DIR.mkdir(parents=True, exist_ok=True)

MIN_SAMPLES_RECOMMENDED = 3
MAX_SAMPLES = 5
MAX_FILE_BYTES = 12 * 1024 * 1024  # 12MB (사진 여유)
ALLOWED_DOC_EXT = {".txt", ".md", ".csv", ".docx", ".xlsx", ".xls", ".pdf", ".json"}
ALLOWED_IMAGE_EXT = {".jpg", ".jpeg", ".png", ".webp", ".gif", ".bmp", ".heic", ".heif"}
ALLOWED_EXT = ALLOWED_DOC_EXT | ALLOWED_IMAGE_EXT

IMAGE_FORM_PROMPT = """당신은 기업 '차량 운행일지' 서식 분석 전문가입니다.
사용자가 올린 사진은 회사 운행일지 양식(또는 작성본)을 촬영한 것입니다.

사진 속 글자·표·칸 구조를 최대한 정확히 읽어, 아래 형식으로 **한국어 텍스트만** 출력하세요.
(마크다운 코드펜스 금지)

[문서 제목]
...

[상단 항목]
- 항목명: 값(있으면)

[표 컬럼]
컬럼1 | 컬럼2 | ...

[표 내용 / 예시 행]
...

[서명·결재란]
...

[레이아웃·서식 특징]
...

읽을 수 없는 부분은 (불명)으로 표기하세요. 추측으로 없는 항목을 만들지 마세요.
"""


def _safe_email(email: str) -> str:
    return "".join(c if c.isalnum() or c in "._-@" else "_" for c in email.strip().lower())


def user_style_dir(email: str) -> Path:
    d = STYLES_DIR / _safe_email(email)
    d.mkdir(parents=True, exist_ok=True)
    (d / "samples").mkdir(exist_ok=True)
    return d


def profile_path(email: str) -> Path:
    return user_style_dir(email) / "profile.json"


def load_profile(email: str) -> dict[str, Any]:
    path = profile_path(email)
    if not path.exists():
        return _empty_profile()
    try:
        with open(path, "r", encoding="utf-8") as f:
            data = json.load(f)
        return {**_empty_profile(), **data}
    except Exception:
        return _empty_profile()


def save_profile(email: str, profile: dict[str, Any]) -> dict[str, Any]:
    path = profile_path(email)
    path.parent.mkdir(parents=True, exist_ok=True)
    with open(path, "w", encoding="utf-8") as f:
        json.dump(profile, f, ensure_ascii=False, indent=2)
    return profile


def _empty_profile() -> dict[str, Any]:
    return {
        "sample_count": 0,
        "samples": [],
        "learned": False,
        "learned_at": None,
        "style_summary": "",
        "tone_rules": [],
        "format_notes": "",
        "purpose_phrases": [],
        "memo_style": "",
        "summary_style": "",
        "place_naming": "",
        "excerpts": [],
        # 회사 서식 구조
        "form_title": "차량 운행일지",
        "form_columns": [],
        "form_header_fields": [],
        "form_layout": "",
        "form_signature": "",
        "active_sample_id": None,
        "engine": None,
        "message": "아직 저장된 회사 서식이 없습니다. 회사 운행일지 양식을 업로드해 주세요.",
    }


def _now() -> str:
    return datetime.now(timezone.utc).isoformat()


# ── 텍스트 추출 ───────────────────────────────────────


def extract_text_from_bytes(filename: str, data: bytes) -> tuple[str, str]:
    """
    파일 바이트 → 텍스트.
    반환: (text, warning)
    사진(jpg/png 등)은 AI 비전으로 서식 분석.
    """
    name = (filename or "file").lower()
    ext = Path(name).suffix.lower()
    if ext not in ALLOWED_EXT:
        raise ValueError(
            f"지원하지 않는 형식입니다: {ext or '(없음)'} "
            "(사진: jpg/png/webp · 문서: pdf/docx/xlsx/txt)"
        )

    if len(data) > MAX_FILE_BYTES:
        raise ValueError("파일이 너무 큽니다. 12MB 이하로 올려 주세요.")

    if ext in ALLOWED_IMAGE_EXT:
        return _extract_from_image(data, ext)

    if ext in {".txt", ".md", ".csv"}:
        text = _decode_bytes(data)
        return text.strip(), ""

    if ext == ".json":
        try:
            obj = json.loads(_decode_bytes(data))
            text = json.dumps(obj, ensure_ascii=False, indent=2)
        except Exception:
            text = _decode_bytes(data)
        return text.strip(), ""

    if ext == ".docx":
        return _extract_docx(data), ""

    if ext in {".xlsx", ".xls"}:
        return _extract_xlsx(data), ""

    if ext == ".pdf":
        text, warn = _extract_pdf(data)
        # 스캔 PDF(텍스트 거의 없음) → 첫 페이지를 이미지처럼 비전 분석은 어려움
        # 사용자에게 사진 촬영 유도
        if len(text) < 40:
            warn = (
                "PDF에서 글자를 거의 읽지 못했습니다. "
                "스캔본이면 양식을 사진으로 찍어 jpg/png로 올려 주세요."
            )
        return text, warn

    raise ValueError(f"지원하지 않는 형식: {ext}")


def _decode_bytes(data: bytes) -> str:
    for enc in ("utf-8", "cp949", "euc-kr", "utf-16"):
        try:
            return data.decode(enc)
        except Exception:
            continue
    return data.decode("utf-8", errors="ignore")


def _extract_docx(data: bytes) -> str:
    import io
    from docx import Document

    doc = Document(io.BytesIO(data))
    parts: list[str] = []
    for p in doc.paragraphs:
        if p.text and p.text.strip():
            parts.append(p.text.strip())
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts).strip()


def _extract_xlsx(data: bytes) -> str:
    import io
    from openpyxl import load_workbook

    wb = load_workbook(io.BytesIO(data), data_only=True, read_only=True)
    parts: list[str] = []
    for ws in wb.worksheets:
        parts.append(f"[시트: {ws.title}]")
        for row in ws.iter_rows(values_only=True):
            cells = [str(c).strip() for c in row if c is not None and str(c).strip()]
            if cells:
                parts.append(" | ".join(cells))
    wb.close()
    return "\n".join(parts).strip()


def _extract_pdf(data: bytes) -> tuple[str, str]:
    import io
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    parts: list[str] = []
    for page in reader.pages:
        t = page.extract_text() or ""
        if t.strip():
            parts.append(t.strip())
    text = "\n".join(parts).strip()
    warn = ""
    if len(text) < 40:
        warn = (
            "PDF에서 텍스트를 거의 읽지 못했습니다. "
            "스캔본이면 양식을 사진으로 찍어 올려 주세요."
        )
    return text, warn


def _image_mime(ext: str) -> str:
    return {
        ".jpg": "image/jpeg",
        ".jpeg": "image/jpeg",
        ".png": "image/png",
        ".webp": "image/webp",
        ".gif": "image/gif",
        ".bmp": "image/bmp",
        ".heic": "image/heic",
        ".heif": "image/heif",
    }.get(ext, "image/jpeg")


def _prepare_image_for_vision(data: bytes, ext: str) -> tuple[bytes, str]:
    """
    비전 API용 이미지 전처리.
    - HEIC → JPEG 변환 시도
    - 过大 해상도/용량 축소
    """
    import io

    mime = _image_mime(ext)
    try:
        from PIL import Image

        # HEIC
        if ext in {".heic", ".heif"}:
            try:
                import pillow_heif  # type: ignore

                pillow_heif.register_heif_opener()
            except Exception:
                raise ValueError(
                    "HEIC 사진은 지원 준비 중입니다. JPG 또는 PNG로 저장해 올려 주세요."
                ) from None

        img = Image.open(io.BytesIO(data))
        if img.mode not in ("RGB", "L"):
            img = img.convert("RGB")
        # 긴 변 2048 이하
        max_side = 2048
        w, h = img.size
        if max(w, h) > max_side:
            scale = max_side / float(max(w, h))
            img = img.resize((int(w * scale), int(h * scale)), Image.Resampling.LANCZOS)

        buf = io.BytesIO()
        img.save(buf, format="JPEG", quality=85, optimize=True)
        return buf.getvalue(), "image/jpeg"
    except ValueError:
        raise
    except Exception:
        # Pillow 실패 시 원본 그대로 (jpeg/png 등)
        if ext in {".heic", ".heif"}:
            raise ValueError(
                "이 사진 형식을 읽지 못했습니다. JPG/PNG로 변환해 올려 주세요."
            ) from None
        return data, mime


def _extract_from_image(data: bytes, ext: str) -> tuple[str, str]:
    """
    사진 속 운행일지 서식을 AI 비전으로 읽어 텍스트화.
    """
    if not OPENAI_API_KEY or OPENAI_API_KEY.startswith("sk-xxxx"):
        raise ValueError(
            "사진 서식 분석에는 OpenAI API 키가 필요합니다. "
            ".env 에 OPENAI_API_KEY 를 설정하거나, 문서를 업로드해 주세요."
        )

    img_bytes, mime = _prepare_image_for_vision(data, ext)
    import base64

    b64 = base64.b64encode(img_bytes).decode("ascii")
    data_url = f"data:{mime};base64,{b64}"

    try:
        from openai import OpenAI

        client = OpenAI(api_key=OPENAI_API_KEY)
        # 비전 지원 모델 (gpt-4o-mini 기본)
        model = OPENAI_MODEL or "gpt-4o-mini"
        resp = client.chat.completions.create(
            model=model,
            temperature=0.1,
            max_tokens=2500,
            messages=[
                {
                    "role": "user",
                    "content": [
                        {"type": "text", "text": IMAGE_FORM_PROMPT},
                        {
                            "type": "image_url",
                            "image_url": {"url": data_url, "detail": "high"},
                        },
                    ],
                }
            ],
        )
        text = (resp.choices[0].message.content or "").strip()
    except Exception as e:
        raise ValueError(f"사진 서식 분석에 실패했습니다: {e}") from e

    if len(text) < 30:
        raise ValueError(
            "사진에서 서식 내용을 충분히 읽지 못했습니다. "
            "밝은 곳에서 양식 전체가 들어오게 다시 찍어 주세요."
        )

    warn = "사진 서식을 AI가 분석해 등록했습니다. 인식이 어긋나면 다른 각도 사진이나 문서를 추가해 주세요."
    return text, warn


# ── 샘플 CRUD ─────────────────────────────────────────


def list_style_status(email: str) -> dict[str, Any]:
    profile = load_profile(email)
    n = len(profile.get("samples") or [])
    active = profile.get("active_sample_id")
    return {
        "sample_count": n,
        "max_samples": MAX_SAMPLES,
        "min_recommended": MIN_SAMPLES_RECOMMENDED,
        "learned": bool(profile.get("learned")),
        "learned_at": profile.get("learned_at"),
        "style_summary": profile.get("style_summary") or "",
        "tone_rules": profile.get("tone_rules") or [],
        "format_notes": profile.get("format_notes") or "",
        "form_title": profile.get("form_title") or "차량 운행일지",
        "form_columns": profile.get("form_columns") or [],
        "form_header_fields": profile.get("form_header_fields") or [],
        "form_layout": profile.get("form_layout") or "",
        "form_signature": profile.get("form_signature") or "",
        "active_sample_id": active,
        "samples": [
            {
                "id": s.get("id"),
                "filename": s.get("filename"),
                "uploaded_at": s.get("uploaded_at"),
                "char_count": s.get("char_count", 0),
                "preview": (s.get("preview") or "")[:180],
                "source_type": s.get("source_type") or "document",
                "is_active": s.get("id") == active or (not active and i == 0),
            }
            for i, s in enumerate(profile.get("samples") or [])
        ],
        "ready": n >= 1,
        "recommended_ready": n >= MIN_SAMPLES_RECOMMENDED,
        "message": _status_message(n, bool(profile.get("learned"))),
        "engine": profile.get("engine"),
    }


def _status_message(n: int, learned: bool) -> str:
    if n == 0:
        return "회사 운행일지 서식(양식)을 업로드하면 저장·학습 후, 그 서식에 맞춰 작성합니다."
    if n < MIN_SAMPLES_RECOMMENDED:
        return f"서식 {n}장 저장됨 · {MIN_SAMPLES_RECOMMENDED - n}장 더 올리면 학습 품질이 좋아집니다. 삭제·추가 가능."
    if learned:
        return f"서식 {n}장 학습 완료 · 생성·인쇄 시 회사 서식·말투가 반영됩니다."
    return f"서식 {n}장 저장됨 · ‘서식 다시 학습’을 눌러 주세요."


def add_sample_from_upload(
    email: str,
    filename: str,
    data: bytes,
) -> dict[str, Any]:
    profile = load_profile(email)
    samples = list(profile.get("samples") or [])
    if len(samples) >= MAX_SAMPLES:
        raise ValueError(f"서식은 최대 {MAX_SAMPLES}장까지 등록할 수 있습니다.")

    text, warn = extract_text_from_bytes(filename, data)
    if len(text.strip()) < 30:
        raise ValueError(
            "서식 내용을 충분히 읽지 못했습니다. "
            "사진을 더 선명하게 찍거나, 문서로 올려 주세요."
            + (f" ({warn})" if warn else "")
        )

    status = _append_sample(email, profile, samples, filename, text, warn)
    # 원본 이미지 저장
    ext = Path(filename or "").suffix.lower()
    if ext in ALLOWED_IMAGE_EXT:
        try:
            sid = (status.get("just_added") or {}).get("id")
            if sid:
                img_path = user_style_dir(email) / "samples" / f"{sid}{ext}"
                img_path.write_bytes(data)
        except Exception:
            pass
    return status


def add_sample_from_text(email: str, title: str, text: str) -> dict[str, Any]:
    profile = load_profile(email)
    samples = list(profile.get("samples") or [])
    if len(samples) >= MAX_SAMPLES:
        raise ValueError(f"샘플은 최대 {MAX_SAMPLES}장까지 등록할 수 있습니다.")
    text = (text or "").strip()
    if len(text) < 30:
        raise ValueError("텍스트가 너무 짧습니다. 실제 운행일지 내용을 붙여넣어 주세요.")
    filename = (title or "붙여넣기 일지").strip() + ".txt"
    return _append_sample(email, profile, samples, filename, text, "")


def _append_sample(
    email: str,
    profile: dict,
    samples: list,
    filename: str,
    text: str,
    warn: str,
) -> dict[str, Any]:
    sid = uuid.uuid4().hex[:12]
    # 원문 저장 (학습용)
    sample_file = user_style_dir(email) / "samples" / f"{sid}.txt"
    sample_file.write_text(text, encoding="utf-8")

    ext = Path(filename or "").suffix.lower()
    is_photo = ext in ALLOWED_IMAGE_EXT
    # 원본 사진 보관 (재분석·확인용)
    if is_photo:
        try:
            img_path = user_style_dir(email) / "samples" / f"{sid}{ext or '.jpg'}"
            # 업로드 시 원본 바이트는 add_sample_from_upload에서 저장
        except Exception:
            pass

    samples.append(
        {
            "id": sid,
            "filename": filename,
            "uploaded_at": _now(),
            "char_count": len(text),
            "preview": text[:200].replace("\n", " "),
            "path": str(sample_file.name),
            "source_type": "photo" if is_photo else "document",
        }
    )
    profile["samples"] = samples
    profile["sample_count"] = len(samples)
    profile["learned"] = False  # 새 샘플 → 재학습 필요
    if not profile.get("active_sample_id"):
        profile["active_sample_id"] = sid
    save_profile(email, profile)

    # 3장 이상이면 자동 학습
    learn_result = None
    if len(samples) >= MIN_SAMPLES_RECOMMENDED:
        learn_result = learn_style(email)
    elif len(samples) >= 1:
        # 1장부터도 간이 학습
        learn_result = learn_style(email)

    status = list_style_status(email)
    status["warning"] = warn or None
    status["just_added"] = {"id": sid, "filename": filename}
    status["learn"] = learn_result
    return status


def delete_sample(email: str, sample_id: str) -> dict[str, Any]:
    """저장된 회사 서식(샘플) 삭제."""
    profile = load_profile(email)
    samples = list(profile.get("samples") or [])
    keep = []
    removed = False
    for s in samples:
        if s.get("id") == sample_id:
            removed = True
            p = user_style_dir(email) / "samples" / f"{sample_id}.txt"
            if p.exists():
                p.unlink()
        else:
            keep.append(s)
    if not removed:
        raise ValueError("해당 서식을 찾을 수 없습니다.")
    profile["samples"] = keep
    profile["sample_count"] = len(keep)
    profile["learned"] = False
    if profile.get("active_sample_id") == sample_id:
        profile["active_sample_id"] = keep[0]["id"] if keep else None
    save_profile(email, profile)
    if keep:
        learn_style(email)
    else:
        save_profile(email, _empty_profile())
    return list_style_status(email)


def set_active_sample(email: str, sample_id: str) -> dict[str, Any]:
    """주 사용 회사 서식 지정 후 재학습."""
    profile = load_profile(email)
    ids = {s.get("id") for s in (profile.get("samples") or [])}
    if sample_id not in ids:
        raise ValueError("해당 서식을 찾을 수 없습니다.")
    profile["active_sample_id"] = sample_id
    profile["learned"] = False
    save_profile(email, profile)
    learn_style(email)
    return list_style_status(email)


def _load_sample_texts(email: str, profile: dict) -> list[tuple[str, str]]:
    """(filename, text) 목록."""
    out = []
    for s in profile.get("samples") or []:
        sid = s.get("id")
        path = user_style_dir(email) / "samples" / f"{sid}.txt"
        if path.exists():
            out.append((s.get("filename") or sid, path.read_text(encoding="utf-8")))
    return out


# ── 스타일 학습 ───────────────────────────────────────


STYLE_ANALYSIS_PROMPT = """당신은 기업 차량 운행일지 서식·문체 분석 전문가입니다.
아래에 한 회사(또는 한 직원)의 실제 운행일지 양식·작성본이 있습니다.

목표:
1) 회사마다 다른 '서식 구조'(제목, 헤더 항목, 표 컬럼, 서명란 등)를 파악
2) 작성 말투·표현 습관을 파악
3) 이후 새 일지를 그 회사 서식에 맞춰, 그 사람이 쓴 것처럼 만들 수 있는 프로필 JSON 생성

반드시 JSON만 출력:
{
  "style_summary": "2~4문장 요약 (말투·문장 길이·격식 정도)",
  "tone_rules": ["규칙1", "규칙2", "규칙3", "규칙4", "규칙5"],
  "format_notes": "표기 습관 (시간, 장소명, 거리, 목적, 비고 쓰는 방식)",
  "purpose_phrases": ["자주 쓰는 목적 표현들"],
  "memo_style": "비고/메모 쓰는 습관",
  "summary_style": "하루 요약 문장 습관",
  "place_naming": "장소 호칭 습관 (본사/사무실/거래처명 등)",
  "do_not": ["이 사람이 절대 안 쓸 법한 표현/과장"],
  "form_title": "문서 제목 (예: 차량 운행일지)",
  "form_header_fields": ["작성일", "차량번호", "운전자", "부서 등 상단 항목들"],
  "form_columns": ["출발시각", "도착시각", "출발지", "도착지", "목적", "거리", "비고 등 표 컬럼"],
  "form_layout": "표/칸 배치, 누적거리·점심·합계 위치 등 레이아웃 설명",
  "form_signature": "작성자/확인자/결재 서명란 유무와 명칭"
}
"""


def learn_style(email: str) -> dict[str, Any]:
    profile = load_profile(email)
    texts = _load_sample_texts(email, profile)
    if not texts:
        raise ValueError("학습할 샘플이 없습니다.")

    # 프롬프트용 발췌 (토큰 절약)
    chunks = []
    for i, (name, text) in enumerate(texts, 1):
        excerpt = text[:3500]
        chunks.append(f"===== 샘플 {i}: {name} =====\n{excerpt}")
    corpus = "\n\n".join(chunks)

    engine = "openai"
    try:
        if OPENAI_API_KEY and not OPENAI_API_KEY.startswith("sk-xxxx"):
            analysis = _analyze_with_openai(corpus)
        else:
            engine = "fallback"
            analysis = _analyze_fallback(texts)
    except Exception as e:
        engine = "fallback"
        analysis = _analyze_fallback(texts)
        analysis["_error"] = str(e)

    excerpts = []
    for name, text in texts[:5]:
        # 의미 있는 중간 발췌
        clean = re.sub(r"\s+", " ", text).strip()
        if len(clean) > 400:
            mid = len(clean) // 4
            snip = clean[mid : mid + 350]
        else:
            snip = clean[:350]
        excerpts.append(f"[{name}] {snip}")

    profile.update(
        {
            "learned": True,
            "learned_at": _now(),
            "style_summary": analysis.get("style_summary") or "",
            "tone_rules": analysis.get("tone_rules") or [],
            "format_notes": analysis.get("format_notes") or "",
            "purpose_phrases": analysis.get("purpose_phrases") or [],
            "memo_style": analysis.get("memo_style") or "",
            "summary_style": analysis.get("summary_style") or "",
            "place_naming": analysis.get("place_naming") or "",
            "do_not": analysis.get("do_not") or [],
            "form_title": analysis.get("form_title") or "차량 운행일지",
            "form_header_fields": analysis.get("form_header_fields") or [],
            "form_columns": analysis.get("form_columns") or [],
            "form_layout": analysis.get("form_layout") or "",
            "form_signature": analysis.get("form_signature") or "",
            "excerpts": excerpts,
            "engine": engine,
            "sample_count": len(texts),
            "message": f"회사 서식 {len(texts)}장 학습 완료",
        }
    )
    if not profile.get("active_sample_id") and profile.get("samples"):
        profile["active_sample_id"] = profile["samples"][0].get("id")
    save_profile(email, profile)
    return {
        "ok": True,
        "engine": engine,
        "sample_count": len(texts),
        "style_summary": profile["style_summary"],
        "tone_rules": profile["tone_rules"],
        "form_title": profile["form_title"],
        "form_columns": profile["form_columns"],
        "learned_at": profile["learned_at"],
    }


def _analyze_with_openai(corpus: str) -> dict[str, Any]:
    from openai import OpenAI

    client = OpenAI(api_key=OPENAI_API_KEY)
    resp = client.chat.completions.create(
        model=OPENAI_MODEL,
        temperature=0.2,
        messages=[
            {"role": "system", "content": STYLE_ANALYSIS_PROMPT},
            {
                "role": "user",
                "content": f"다음 운행일지 샘플들을 분석하세요.\n\n{corpus[:14000]}",
            },
        ],
        response_format={"type": "json_object"},
    )
    content = resp.choices[0].message.content or "{}"
    return json.loads(content)


def _analyze_fallback(texts: list[tuple[str, str]]) -> dict[str, Any]:
    """키 없을 때 휴리스틱 분석."""
    joined = "\n".join(t for _, t in texts)
    purposes = re.findall(
        r"(업무\s*출장|고객\s*미팅|거래처\s*방문|납품|배송|중식|업무\s*복귀|현장\s*점검|회의)",
        joined,
    )
    purpose_phrases = list(dict.fromkeys(purposes))[:8] or ["업무 출장", "거래처 방문", "업무 복귀"]

    # 문장 길이 추정
    sentences = re.split(r"[.\n]", joined)
    sentences = [s.strip() for s in sentences if len(s.strip()) > 4]
    avg_len = sum(len(s) for s in sentences) / max(1, len(sentences))
    if avg_len < 18:
        tone = "짧고 개조식에 가까운 간결체"
    elif avg_len < 35:
        tone = "보통 길이의 실무 보고체"
    else:
        tone = "서술형에 가까운 정중한 보고체"

    has_honorific = any(k in joined for k in ("하였", "했습니다", "드렸", "방문함", "실시"))
    formal = "합니다/하였음 계열의 격식체" if has_honorific or "함" in joined else "간결한 실무 단문"

    return {
        "style_summary": f"회사 서식 {len(texts)}장 기준. {tone}. {formal}를 사용하며 제출용 운행일지 톤을 유지합니다.",
        "tone_rules": [
            "과장을 피하고 사실 위주로 기술",
            "장소·목적·거리는 간결하게",
            formal,
            "하루 요약은 한 문장으로 마무리",
            "비고는 필요할 때만 짧게",
        ],
        "format_notes": "시간은 HH:MM, 거리는 km, 목적·비고는 짧은 명사구/개조식 선호",
        "purpose_phrases": purpose_phrases,
        "memo_style": "필수 정보만 짧게, 수식어 최소화",
        "summary_style": "방문지 중심 한 줄 요약",
        "place_naming": "본사·거래처·고객사 등 실무 호칭 유지",
        "do_not": ["이모지", "구어체 반말", "불필요한 수식어", "소설형 서술"],
        "form_title": "차량 운행일지",
        "form_header_fields": ["작성일", "차량번호", "운전자", "회사명", "최초누적", "종료누적"],
        "form_columns": ["출발시각", "도착시각", "출발지", "도착지", "운행목적", "거리(km)", "비고"],
        "form_layout": "상단 메타 정보 + 구간 표 + 하단 요약/합계",
        "form_signature": "작성자 · 확인자 서명란",
    }


def build_style_prompt_block(email: str) -> str:
    """생성 프롬프트에 넣을 회사 서식 + 말투 블록."""
    profile = load_profile(email)
    if not profile.get("samples"):
        return ""

    rules = profile.get("tone_rules") or []
    rules_txt = "\n".join(f"- {r}" for r in rules) if rules else "- (규칙 없음)"
    purposes = ", ".join(profile.get("purpose_phrases") or []) or "(없음)"
    do_not = ", ".join(profile.get("do_not") or []) or "(없음)"
    cols = ", ".join(profile.get("form_columns") or []) or "(기본 구간 표)"
    headers = ", ".join(profile.get("form_header_fields") or []) or "(기본 헤더)"
    excerpts = profile.get("excerpts") or []
    excerpt_block = "\n\n".join(excerpts[:3]) if excerpts else "(발췌 없음)"

    return f"""
[회사 전용 운행일지 서식 + 사용자 말투 — 반드시 따를 것]
※ 사용자가 업로드한 실제 회사 운행일지 양식·작성본에서 학습한 내용입니다.
※ 새 일지는 이 회사 서식 구조에 맞추고, 사용자가 직접 쓴 것처럼 보여야 합니다.
※ 일반적인 AI 문체·임의 양식을 쓰지 마세요.

■ 서식 구조
- 문서 제목: {profile.get('form_title') or '차량 운행일지'}
- 상단 항목: {headers}
- 표 컬럼: {cols}
- 레이아웃: {profile.get('form_layout') or '(표준)'}
- 서명/결재: {profile.get('form_signature') or '(작성자·확인자)'}
- 표기 습관: {profile.get('format_notes') or '(미상)'}

■ 말투
- 요약: {profile.get('style_summary') or '(미학습)'}
- 목적 표현: {purposes}
- 비고 습관: {profile.get('memo_style') or '(미상)'}
- 요약 습관: {profile.get('summary_style') or '(미상)'}
- 장소 호칭: {profile.get('place_naming') or '(미상)'}
- 말투 규칙:
{rules_txt}
- 피할 표현: {do_not}

[실제 회사 서식/작성 발췌 — 어휘·컬럼·호흡 참고]
{excerpt_block}
"""

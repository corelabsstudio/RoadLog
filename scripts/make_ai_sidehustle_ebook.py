# -*- coding: utf-8 -*-
"""AI 부업 준비 전자책 초안 + 부록을 바탕화면 크몽_전자책 폴더에 생성."""
from __future__ import annotations

from pathlib import Path
import time

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

ROOT = Path.home() / "Desktop" / "크몽_전자책"
DIR_MS = ROOT / "01_원고"
DIR_TPL = ROOT / "02_부록_템플릿"
DIR_CHK = ROOT / "03_체크리스트"
DIR_COVER = ROOT / "04_표지_안내"
for d in (ROOT, DIR_MS, DIR_TPL, DIR_CHK, DIR_COVER):
    d.mkdir(parents=True, exist_ok=True)

DOCX = DIR_MS / "AI로_부업_준비하기_전자책_원고.docx"
HWP = DIR_MS / "AI로_부업_준비하기_전자책_원고.hwp"


def build_docx() -> Path:
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = Cm(2)
        sec.bottom_margin = Cm(2)
        sec.left_margin = Cm(2.5)
        sec.right_margin = Cm(2.5)

    def set_run_font(run, size=11, bold=False):
        run.font.size = Pt(size)
        run.font.name = "맑은 고딕"
        run.bold = bold
        try:
            run._element.rPr.rFonts.set(
                "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia",
                "맑은 고딕",
            )
        except Exception:
            pass

    def h1(t):
        p = doc.add_heading(t, level=1)
        return p

    def h2(t):
        return doc.add_heading(t, level=2)

    def h3(t):
        return doc.add_heading(t, level=3)

    def body(t):
        p = doc.add_paragraph()
        r = p.add_run(t)
        set_run_font(r, 11)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.line_spacing = 1.4
        return p

    def bullet(t):
        p = doc.add_paragraph(t, style="List Bullet")
        for r in p.runs:
            set_run_font(r, 11)
        return p

    def num(t):
        p = doc.add_paragraph(t, style="List Number")
        for r in p.runs:
            set_run_font(r, 11)
        return p

    def center(t, size=12, bold=False):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(t)
        set_run_font(r, size, bold)
        return p

    # ----- 표지 -----
    center("\n\n", 12)
    center("개발 몰라도 따라 하는", 14)
    center("AI로 부업 준비하기", 26, True)
    center("\n실전 노트 · 템플릿 · 체크리스트", 14)
    center("\n과장 없는 과정 중심 전자책 초안", 12)
    center("\n\nONHAE", 12)
    center("버전 0.1 초안", 10)
    body("")
    body("읽는 분께: 이 책은 이미 큰돈을 벌었다는 성공 자랑기가 아닙니다. AI를 써서 부업 상품과 판매 준비를 해 본 과정을, 쉬운 말로 정리한 실습 노트입니다. 수익은 실행과 시장에 따라 달라지며 보장되지 않습니다.")
    doc.add_page_break()

    # ----- 목차 -----
    h1("목차")
    for line in [
        "머리말 — 이 책을 쓰는 이유",
        "1장  이 책의 약속과 한계",
        "2장  AI란 무엇인가 (아주 쉽게)",
        "3장  컴맹은 아니지만, 개발은 모르는 분 기준",
        "4장  부업 아이디어를 AI로 고르는 법",
        "5장  팔 수 있는 형태로 다듬기",
        "6장  판매 글·FAQ·가격을 AI로 쓰기",
        "7장  포트폴리오와 이미지를 준비하기",
        "8장  크몽·숨고에 올려 보는 흐름",
        "9장  하면 안 되는 것들",
        "10장 일주일 실행 계획",
        "11장 사례 노트 (준비 과정 기록)",
        "12장 자주 묻는 질문",
        "맺음말",
        "부록 안내",
    ]:
        body(line)
    doc.add_page_break()

    # ----- 머리말 -----
    h1("머리말 — 이 책을 쓰는 이유")
    body("AI 이야기가 어디에나 나옵니다. 유튜브와 전자책에는 금방 돈을 벌 수 있다는 말도 많습니다. 그런데 막상 시작하려면 어디서부터 눌러야 할지 막막합니다.")
    body("저는 개발자가 아니어도 할 수 있는 일부터, AI에게 물어 가며 하나씩 준비해 보았습니다. 서비스 설명 쓰기, 자주 묻는 질문 만들기, 가격 나누기, 포트폴리오 구성하기, 프로필 사진과 소개 문구 정하기 같은 일입니다.")
    body("그 과정에서 느낀 점은 분명합니다. AI는 마법이 아닙니다. 대신 혼자 白紙에서 고민하는 시간을 많이 줄여 줍니다. 이 책은 그 줄이는 방법을, 어려운 용어 없이 순서대로 적었습니다.")
    body("아직 큰 수익을 인증할 단계는 아닙니다. 그래서 이 책은 수익 자랑 대신, 준비 과정에서 바로 써먹을 문장과 체크리스트를 드립니다. 같은 출발선에 있는 분에게 더 도움이 되기를 바랍니다.")
    doc.add_page_break()

    # ----- 1 -----
    h1("1장  이 책의 약속과 한계")
    h2("1.1 이 책이 약속하는 것")
    bullet("AI 채팅을 부업 준비에 쓰는 구체적인 순서")
    bullet("복사해서 고칠 수 있는 문장 템플릿")
    bullet("한 주 동안 따라 할 수 있는 할 일 표")
    bullet("과장 광고에 속지 않기 위한 기준")

    h2("1.2 이 책이 약속하지 않는 것")
    bullet("특정 기간 안에 무조건 수익")
    bullet("하루 만에 부자 되는 비법")
    bullet("불법·편법·스팸으로 돈 버는 방법")
    bullet("코딩 전문가가 되는 완전한 강의")

    h2("1.3 추천 독자")
    body("마우스와 키보드로 파일을 저장하고, 인터넷 검색과 회원가입 정도는 할 수 있는 분. 스마트폰이나 PC로 ChatGPT 같은 AI 화면에 글을 입력해 본 적이 있거나, 앞으로 해 볼 의지가 있는 분.")
    body("완전 처음 전원을 켜는 단계라면, 먼저 주변 또는 기초 컴퓨터 도움으로 파일·폴더·회원가입만 익힌 뒤 이 책을 보시면 좋습니다.")

    h2("1.4 읽기 전에 준비할 것")
    num("인터넷이 되는 컴퓨터 또는 노트북")
    num("이메일 계정")
    num("AI 채팅 도구 하나 (무료 요금제부터 가능)")
    num("메모장 또는 한글·워드 아무거나")
    num("노트 앱이나 종이에 적을 준비")
    doc.add_page_break()

    # ----- 2 -----
    h1("2장  AI란 무엇인가 (아주 쉽게)")
    h2("2.1 한 문장으로")
    body("AI 채팅은, 내가 글이나 질문을 넣으면 그에 맞는 글·아이디어·초안을 만들어 주는 도우미입니다. 검색과 비슷하지만, 여러 문장을 이어서 다듬어 달라고 부탁할 수 있습니다.")

    h2("2.2 잘 하는 일")
    bullet("긴 설명을 짧게 줄이기")
    bullet("판매 글·FAQ 초안 쓰기")
    bullet("표나 단계 나누기")
    bullet("어려운 말을 쉬운 말로 바꾸기")
    bullet("빠진 항목 지적해 달라고 하기")

    h2("2.3 못 믿으면 안 되는 일")
    bullet("최신 법률·세금·의료를 최종 결정으로 믿기")
    bullet("확인 없이 숫자·사례를 사실이라고 쓰기")
    bullet("내 개인정보·비밀번호를 그대로 붙여 넣기")
    bullet("틀린 코드를 검토 없이 실행하기")

    h2("2.4 기본 사용법 5단계")
    num("하고 싶은 일을 한 줄로 쓴다. 예: 엑셀 자동화 서비스 소개 문단 써줘.")
    num("대상 독자를 적는다. 예: 개발 모르는 직장인.")
    num("길이·말투를 정한다. 예: 200자, 친절한 말투, 과한 광고 말투 금지.")
    num("나온 글을 읽으며 사실과 다른 부분을 고친다.")
    num("다시 부탁한다. 예: 더 쉽게, 반말로 하지 말고, 리스트로.")

    h2("2.5 복붙용 기본 질문 틀")
    body("아래를 복사해 빈칸만 채우면 됩니다.")
    body("역할: 너는 초보도 이해하게 설명하는 도우미다.")
    body("목표: (하고 싶은 일)")
    body("독자: (누구를 위한 글인지)")
    body("조건: 쉬운 말, 과장 금지, 단계로 나눠 줘.")
    body("출력: (문단 / 표 / 체크리스트 중 무엇)")
    doc.add_page_break()

    # ----- 3 -----
    h1("3장  개발은 모르는 분 기준")
    h2("3.1 이 책에서 말하는 쉬움의 기준")
    body("클릭, 타이핑, 파일 저장, 회원가입, 결제 창 열기 정도는 가능하다고 가정합니다. 프로그래밍 언어를 외울 필요는 없습니다.")

    h2("3.2 어려우면 여기서 멈추라는 신호")
    bullet("안내와 전혀 다른 화면이 나와 30분 이상 혼자 헤맬 때 → 잠시 멈추고 캡처를 저장한 뒤, 다음 날 다시 하거나 주변에 물어보세요.")
    bullet("결제·개인정보 창이 의심스러울 때 → 진행하지 마세요.")

    h2("3.3 폴더 하나 만들기")
    body("바탕화면에 부업준비 폴더를 만드세요. 그 안에 글초안, 이미지, 체크리스트 하위 폴더를 두면 나중에 헤매지 않습니다.")
    doc.add_page_break()

    # ----- 4 -----
    h1("4장  부업 아이디어를 AI로 고르는 법")
    h2("4.1 세 가지로 거르기")
    body("아이디어가 많아 보이면 오히려 아무것도 못 합니다. 아래 세 가지를 통과한 것만 남기세요.")
    num("내가 이미 조금 해 본 일 또는 배울 의지가 있는 일인가")
    num("다른 사람이 돈·시간을 아끼고 싶어 하는 일인가")
    num("불법·사기·스팸에 가깝지 않은가")

    h2("4.2 AI에게 물어볼 질문 예시")
    body("나는 Windows를 쓰고, 엑셀을 조금 다룰 수 있다. 개발은 모른다. 원격으로 할 수 있는 부업 아이디어 10개를, 난이도 하·중으로만 짧게 나열해 줘. 각각 한 줄로 왜 돈이 될 수 있는지 적어 줘. 과장 없이.")

    h2("4.3 3개로 줄이기")
    body("10개 중 마음에 드는 3개만 고른 뒤, 각각에 대해 다시 물어보세요.")
    body("이 아이디어로 첫 상품을 만든다면, 팔 단위(한 건)를 무엇으로 정하면 좋을까? 초보 기준으로 범위를 좁혀 줘.")

    h2("4.4 첫 상품은 작게")
    body("처음부터 앱을 만들겠다는 목표는 대부분 너무 큽니다. 첫 상품은 작게: 설명 글 작성 대행, 템플릿 한 종, 엑셀 정리 한 가지, 체크리스트 전자책처럼 끝낼 수 있는 단위가 좋습니다.")
    doc.add_page_break()

    # ----- 5 -----
    h1("5장  팔 수 있는 형태로 다듬기")
    h2("5.1 상품의 한 줄 정의")
    body("빈칸을 채우세요.")
    body("나는 (누구) 에게 (결과) 를 (방법) 으로 제공한다.")
    body("예: 나는 엑셀 마감이 반복되는 사람에게, 시트 합치기 결과를, 원격 맞춤 작업으로 제공한다.")

    h2("5.2 포함 / 미포함을 미리 적기")
    body("나중에 수정 분쟁을 줄입니다. AI에게 이렇게 부탁하세요.")
    body("다음 상품의 포함 항목 5개와 미포함 항목 5개를 적어 줘. 초보 판매자 기준으로 과대 약속 금지. 상품: (한 줄 정의)")

    h2("5.3 가격을 나누는 간단한 방법")
    body("한 가지만 하는 기본, 두 가지 연계, 복잡한 맞춤 세 단계로 나누면 구매자가 비교하기 쉽습니다. 시장 시세는 같은 플랫폼에서 비슷한 상품 제목과 가격을 여러 개 살펴본 뒤, 너무 싸지도 비싸지도 않은 가운데를 고르세요.")
    body("처음에는 기본 가격을 너무 낮추지 않는 편이 낫습니다. 수정 요청이 생기면 시간이 곧 비용입니다.")
    doc.add_page_break()

    # ----- 6 -----
    h1("6장  판매 글·FAQ·가격 설명을 AI로 쓰기")
    h2("6.1 서비스 설명 초안 요청문")
    body("다음 조건으로 서비스 상세 설명을 써 줘. 200자 이상. 쉬운 말. 과장 금지. 특수문자 중 작은따옴표와 꺾쇠는 쓰지 마. 상품: (한 줄 정의). 포함: (목록). 진행 순서: 문의-범위확인-제작-납품.")

    h2("6.2 FAQ 요청문")
    body("초보 구매자가 물을 법한 질문 10개와 짧은 답변을 만들어 줘. 답변은 3~5문장 이내. 가격·기간·수정 범위·준비물을 반드시 포함.")

    h2("6.3 제목이 짧아야 할 때")
    body("플랫폼마다 제목 글자 수 제한이 있습니다. 예: 30자 이내. AI에게 반드시 글자 수 제한을 말하세요.")
    body("다음 의미를 유지한 채 제목 후보 8개를 30자 이내로 만들어 줘. 의미: 엑셀 업무 자동화 맞춤 제작.")

    h2("6.4 반드시 사람이 고칠 부분")
    bullet("내가 실제로 못 하는 일 삭제")
    bullet("기간·수정 횟수를 내 일정에 맞게 수정")
    bullet("가짜 경력·가짜 후기 넣지 않기")
    doc.add_page_break()

    # ----- 7 -----
    h1("7장  포트폴리오와 이미지 준비")
    h2("7.1 포트폴리오가 중요한 이유")
    body("처음 보는 사람은 당신을 모릅니다. 그래서 전에 만든 결과물 샘플이 있으면 신뢰가 올라갑니다. 수익 인증이 없어도, 샘플 결과물은 준비할 수 있습니다.")

    h2("7.2 샘플을 만드는 쉬운 방법")
    num("가상 데이터로 예제 표 만들기 (진짜 회사 비밀 데이터 금지)")
    num("작업 전 모습과 작업 후 모습을 글로 설명")
    num("AI에게 포트폴리오 제목·한 줄 요약 부탁")
    num("이미지 규격은 플랫폼 안내를 그대로 따르기")

    h2("7.3 이미지 업로드가 안 될 때")
    body("규격·용량·브라우저·네트워크 문제일 수 있습니다. 다른 브라우저, 시크릿 창, 간단한 테스트 이미지로 먼저 확인하세요. 임시 저장 후 나중에 다시 시도해도 됩니다.")
    doc.add_page_break()

    # ----- 8 -----
    h1("8장  크몽·숨고에 올려 보는 흐름")
    h2("8.1 두 플랫폼의 차이 (쉽게)")
    body("크몽은 상품을 미리 올려 두고 구매자가 찾아오는 쪽에 가깝습니다. 숨고는 손님이 요청을 올리면 고수가 견적을 보내는 방식에 가깝습니다. 숨고 견적은 캐시가 필요할 수 있으니, 될 만한 요청만 고르세요.")

    h2("8.2 등록 전 체크")
    bullet("증명할 수 없는 경력·학력 문구 삭제")
    bullet("서비스 범위와 미포함 명시")
    bullet("수정 규정 작성")
    bullet("연락·작업은 플랫폼 안에서")

    h2("8.3 승인 전후")
    body("승인이 나기 전에는 보통 일반 구매 노출이 어렵습니다. 승인 알림을 기다리고, 반려 사유가 오면 그 부분만 고치면 됩니다.")

    h2("8.4 첫 문의가 오면")
    num("무엇을 원하는지 한 줄로 다시 확인")
    num("파일·기한·수정 범위 확인")
    num("가능하면 작은 범위로 먼저 끝내기")
    num("추가 요청은 추가 견적으로 안내")
    doc.add_page_break()

    # ----- 9 -----
    h1("9장  하면 안 되는 것들")
    bullet("허위 수익·허위 후기")
    bullet("무단 자동 도배, 스팸성 홍보")
    bullet("계정·인증 우회, 불법 프로그램 안내")
    bullet("과도한 개인정보를 AI에 붙여 넣기")
    bullet("할 수 없는 일을 할 수 있다고 쓰기")
    body("짧게 돈을 벌려다 계정과 신뢰를 잃는 경우가 많습니다. 부업은 오래 가는 쪽이 이득입니다.")
    doc.add_page_break()

    # ----- 10 -----
    h1("10장  일주일 실행 계획")
    h2("Day 1")
    body("아이디어 10개 받기 → 3개로 줄이기 → 한 줄 정의 쓰기")
    h2("Day 2")
    body("포함·미포함·가격 3단계 초안")
    h2("Day 3")
    body("서비스 설명·FAQ 초안 AI로 작성 후 직접 수정")
    h2("Day 4")
    body("포트폴리오 샘플 1~2개 (가상 데이터)")
    h2("Day 5")
    body("프로필 사진·한 줄 소개·활동명 정하기")
    h2("Day 6")
    body("플랫폼 등록 또는 임시 저장, 막히는 부분 메모")
    h2("Day 7")
    body("전체 글 다시 읽기, 과장 표현 삭제, 다음 주 할 일 3개만 정하기")
    body("하루를 못 지키면 일정을 늘리면 됩니다. 중요한 것은 완벽한 한 주가 아니라, 멈춘 곳을 다시 시작하는 것입니다.")
    doc.add_page_break()

    # ----- 11 -----
    h1("11장  사례 노트 (준비 과정 기준)")
    h2("11.1 이 장의 성격")
    body("아래는 매출 자랑 사례가 아닙니다. 실제로 준비하며 마주친 일과, 그때 쓰면 좋았던 판단 기준을 정리한 노트입니다. 독자 여러분도 자기만의 과정 일지로 바꿔 쓰세요.")

    h2("11.2 사례 A — 상품 문장 만들기")
    body("상황: 무엇을 파는지 한 페이지로 써야 했다.")
    body("AI 활용: 초안 생성 → 글자 수 제한에 맞게 다시 요청 → 내가 못 하는 약속 삭제.")
    body("결과: 제목, 상세 설명, FAQ, 절차, 준비물 문장을 확보했다.")
    body("교훈: 첫 초안을 그대로 올리지 말고, 반드시 내가 할 수 있는 범위로 줄인다.")

    h2("11.3 사례 B — 가격 나누기")
    body("상황: 기본·표준·고급 가격을 정해야 했다.")
    body("방법: 비슷한 상품 가격을 여러 개 살펴보고, 너무 낮은 덤핑과 너무 높은 가격을 피했다.")
    body("교훈: 싼 가격은 주문을 부를 수도 있지만, 수정 부담이 커지면 손해가 된다.")

    h2("11.4 사례 C — 포트폴리오 샘플")
    body("상황: 실거래 전이라 보여줄 고객 결과물이 부족했다.")
    body("방법: 가상 데이터로 작업 전후를 만들고, 처리 단계를 카드 이미지로 정리했다.")
    body("교훈: 수익 인증이 없어도 샘플 결과물은 만들 수 있다. 다만 실제 고객 비밀 데이터는 쓰지 않는다.")

    h2("11.5 사례 D — 업로드·등록 장애")
    body("상황: 이미지가 규격에 맞아 보여도 업로드가 실패했다.")
    body("대응: 임시 저장, 다른 파일·브라우저 시도, 원인 메모.")
    body("교훈: 기술 문제와 내용 문제를 구분한다. 규격만 고치면 되는 일과, 플랫폼·네트워크 문제를 나누면 마음이 덜 지친다.")

    h2("11.6 사례 E — 채널 두 개")
    body("상황: 한곳은 상품형, 다른 한곳은 견적형 플랫폼이었다.")
    body("대응: 글 전체를 복붙하지 않고, 견적용 짧은 문구로 바꿨다. 견적 비용이 드는 쪽은 선별해서 넣기로 했다.")
    body("교훈: 채널마다 문장 길이와 비용 구조가 다르다.")
    doc.add_page_break()

    # ----- 12 -----
    h1("12장  자주 묻는 질문")
    h3("Q1. AI만 있으면 돈이 되나요?")
    body("A. AI만으로 자동 입금되지 않습니다. 상품·신뢰·실행이 필요합니다. AI는 준비 속도를 높여 줍니다.")
    h3("Q2. 정말 개발을 몰라도 되나요?")
    body("A. 이 책 범위의 글 쓰기·등록·템플릿 활용은 가능합니다. 프로그램을 대신 만들어 주는 일까지 하려면 배우거나 외주·협업이 필요합니다.")
    h3("Q3. 무료 AI로 충분한가요?")
    body("A. 시작은 무료로 가능합니다. 사용량·속도 제한이 있으면 유료를 나중에 검토하면 됩니다.")
    h3("Q4. 얼마나 걸려요?")
    body("A. 하루 1~2시간 기준 1~2주면 기본 등록 준비까지 가 볼 수 있는 경우가 많습니다. 개인 차이가 큽니다.")
    h3("Q5. 수익이 없으면 책을 사도 소용 없나요?")
    body("A. 이 책은 수익 대체가 아니라 준비 도구입니다. 실행하지 않으면 어떤 책도 소용 없습니다.")
    doc.add_page_break()

    # ----- 맺음말 -----
    h1("맺음말")
    body("부업은 거창한 한 방이 아니라, 작은 상품을 끝까지 올리는 일의 반복에 가깝습니다. AI는 그 반복에서 글을 다듬고 빈칸을 채워 주는 도구입니다.")
    body("오늘 할 일이 막막하면, 딱 세 가지만 하세요. 한 줄 정의 쓰기, 포함·미포함 다섯 줄 쓰기, AI로 소개 문단 한 개 만들기. 그것만으로도 어제보다 한 걸음입니다.")
    body("이 초안을 읽으시며 어색한 문장, 더 쉬운 표현, 추가하고 싶은 사례가 있으면 표시해 두세요. 다음 수정본에 반영할 수 있습니다.")
    body("감사합니다.")
    body("ONHAE · AI로 부업 준비하기 · 초안 0.1")

    doc.save(DOCX)
    print("DOCX", DOCX, DOCX.stat().st_size)
    return DOCX


def convert_hwp(docx_path: Path, hwp_path: Path) -> bool:
    import pythoncom
    import win32com.client

    pythoncom.CoInitialize()
    hwp = None
    try:
        hwp = win32com.client.Dispatch("HWPFrame.HwpObject")
        for mod in ("FilePathCheckerModule", "FilePathCheckerModuleExample"):
            try:
                hwp.RegisterModule("FilePathCheckDLL", mod)
                break
            except Exception:
                pass
        try:
            hwp.XHwpWindows.Item(0).Visible = False
        except Exception:
            pass

        # new + insert text from docx paragraphs (reliable enough for draft)
        hwp.HAction.Run("FileNew")
        d = Document(str(docx_path))
        chunks: list[str] = []
        buf = []
        for p in d.paragraphs:
            buf.append(p.text)
            if sum(len(x) for x in buf) > 2500:
                chunks.append("\n".join(buf))
                buf = []
        if buf:
            chunks.append("\n".join(buf))

        for part in chunks:
            hwp.HAction.GetDefault("InsertText", hwp.HParameterSet.HInsertText.HSet)
            hwp.HParameterSet.HInsertText.Text = part + "\n"
            hwp.HAction.Execute("InsertText", hwp.HParameterSet.HInsertText.HSet)

        if hwp_path.exists():
            hwp_path.unlink()
        save_path = str(hwp_path.resolve())
        try:
            hwp.HAction.GetDefault("FileSaveAs_S", hwp.HParameterSet.HFileOpenSave.HSet)
            hwp.HParameterSet.HFileOpenSave.filename = save_path
            hwp.HParameterSet.HFileOpenSave.Format = "HWP"
            hwp.HAction.Execute("FileSaveAs_S", hwp.HParameterSet.HFileOpenSave.HSet)
        except Exception as e:
            print("save err", e)
            return False
        time.sleep(0.8)
        ok = hwp_path.exists()
        print("HWP", hwp_path, ok, hwp_path.stat().st_size if ok else 0)
        return ok
    finally:
        if hwp is not None:
            try:
                hwp.Quit()
            except Exception:
                pass
        pythoncom.CoUninitialize()


def write_support_files() -> None:
    (ROOT / "README.txt").write_text(
        """크몽 전자책 폴더 안내
====================

제목(가제): 개발 몰라도 따라 하는 AI로 부업 준비하기
버전: 0.1 초안
목적: 크몽 전자책 판매용 원고 · 부록 모음

폴더
- 01_원고 : 본문 (HWP, DOCX)
- 02_부록_템플릿 : 복사해서 쓰는 문장
- 03_체크리스트 : 실행 체크
- 04_표지_안내 : 표지 만들 때 참고

읽는 순서
1) README
2) 01_원고 HWP 또는 DOCX
3) 부록 템플릿으로 실습

수정 요청
- 어색한 문장, 더 쉬운 말, 추가 사례를 메모해 주세요.

주의
- 수익 보장 표현 없음
- 허위 후기·불법 자동화 안내 없음
""",
        encoding="utf-8",
    )

    (DIR_TPL / "01_한줄_정의_템플릿.txt").write_text(
        "나는 (누구) 에게 (결과) 를 (방법) 으로 제공한다.\n\n예)\n나는 엑셀 마감이 반복되는 사람에게, 시트를 합친 결과물을, 원격 작업으로 제공한다.\n",
        encoding="utf-8",
    )
    (DIR_TPL / "02_AI_질문_기본틀.txt").write_text(
        "역할: 너는 초보도 이해하게 설명하는 도우미다.\n"
        "목표: (하고 싶은 일)\n"
        "독자: (누구를 위한 글인지)\n"
        "조건: 쉬운 말, 과장 금지, 단계로 나눠 줘.\n"
        "출력: (문단 / 목록 / 표 중 선택)\n",
        encoding="utf-8",
    )
    (DIR_TPL / "03_서비스_설명_요청문.txt").write_text(
        "다음 조건으로 서비스 상세 설명을 써 줘.\n"
        "- 200자 이상\n"
        "- 쉬운 말\n"
        "- 과장 금지\n"
        "- 작은따옴표와 꺾쇠 문자 사용 금지\n"
        "상품 한 줄: (작성)\n"
        "포함: (작성)\n"
        "미포함: (작성)\n"
        "진행 순서: 문의 → 범위 확인 → 작업 → 납품\n",
        encoding="utf-8",
    )
    (DIR_TPL / "04_FAQ_요청문.txt").write_text(
        "초보 구매자가 물을 법한 질문 10개와 짧은 답변을 만들어 줘.\n"
        "답변은 3~5문장 이내.\n"
        "가격, 기간, 수정 범위, 준비물을 반드시 포함.\n"
        "상품: (한 줄 정의)\n",
        encoding="utf-8",
    )
    (DIR_TPL / "05_숨고_견적_템플릿.txt").write_text(
        "안녕하세요. (서비스) 원격/출장 가능 범위를 안내드립니다.\n\n"
        "요청하신 내용 기준으로 (할 일) 가능합니다.\n\n"
        "필요 자료: (파일·증상·기한)\n"
        "대략 일정: ( )\n"
        "견적: 범위 확인 후 안내 (단순 기준 약 00만 원대)\n\n"
        "확인해 보시고 편히 답 주세요.\n",
        encoding="utf-8",
    )
    (DIR_TPL / "06_포함_미포함_표.txt").write_text(
        "[포함]\n1.\n2.\n3.\n4.\n5.\n\n[미포함]\n1.\n2.\n3.\n4.\n5.\n",
        encoding="utf-8",
    )

    (DIR_CHK / "일주일_실행_체크리스트.txt").write_text(
        "Day1 아이디어 10→3, 한 줄 정의\n"
        "Day2 포함/미포함, 가격 3단계\n"
        "Day3 상세 설명·FAQ 초안\n"
        "Day4 포트폴리오 샘플\n"
        "Day5 프로필·한 줄 소개\n"
        "Day6 플랫폼 등록/임시저장\n"
        "Day7 과장 표현 삭제, 다음 주 할 일 3개\n",
        encoding="utf-8",
    )
    (DIR_CHK / "등록_전_체크리스트.txt").write_text(
        "[ ] 증명 못 하는 경력·학력 문구 삭제\n"
        "[ ] 포함/미포함 명시\n"
        "[ ] 수정 규정 작성\n"
        "[ ] 제목 글자 수 확인\n"
        "[ ] 이미지 규격 확인\n"
        "[ ] 과장·허위 수익 표현 없음\n"
        "[ ] 임시저장 또는 제출\n",
        encoding="utf-8",
    )
    (DIR_CHK / "첫문의_응대_체크리스트.txt").write_text(
        "[ ] 원하는 결과 한 줄 확인\n"
        "[ ] 파일/증상/기한 확인\n"
        "[ ] 내 가능 범위인지 판단\n"
        "[ ] 견적·기간 안내\n"
        "[ ] 추가 요청은 추가 비용 고지\n",
        encoding="utf-8",
    )

    (DIR_COVER / "표지_문구_후보.txt").write_text(
        "메인 제목:\n개발 몰라도 따라 하는\nAI로 부업 준비하기\n\n"
        "부제:\n실전 노트 · 템플릿 · 체크리스트\n과장 없는 과정 중심\n\n"
        "하단:\nONHAE\n수익 보장 없음 · 준비와 실행 도구\n",
        encoding="utf-8",
    )
    (DIR_COVER / "상세페이지_판매글_초안.txt").write_text(
        "개발을 몰라도, AI로 부업 준비를 시작해 보고 싶은 분을 위한 실전 노트입니다.\n\n"
        "포함: 쉬운 설명, 따라 하는 순서, 복사해 쓰는 템플릿, 일주일 체크리스트\n"
        "이 책은 이미 큰 수익을 인증하는 성공담이 아닙니다. 준비 과정을 정리한 실행 도구입니다.\n"
        "수익은 보장되지 않으며, 실행량에 따라 달라집니다.\n",
        encoding="utf-8",
    )
    print("support files ok")


def main() -> None:
    write_support_files()
    build_docx()
    ok = convert_hwp(DOCX, HWP)
    # list tree
    print("ROOT", ROOT)
    for p in sorted(ROOT.rglob("*")):
        if p.is_file():
            print(p.relative_to(ROOT), p.stat().st_size)
    if not ok:
        print("HWP conversion failed; DOCX is available")


if __name__ == "__main__":
    main()

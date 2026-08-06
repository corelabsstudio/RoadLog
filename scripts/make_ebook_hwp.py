# -*- coding: utf-8 -*-
"""Create ebook draft DOCX then convert to HWP via Hancom COM."""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import win32com.client
import pythoncom
import time

desk = Path.home() / "Desktop"
docx_path = desk / "엑셀_파이썬_자동화_전자책_원고.docx"
hwp_path = desk / "엑셀_파이썬_자동화_전자책_원고.hwp"

doc = Document()
# narrow-ish for more pages feel but readable
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

def h1(text):
    p = doc.add_heading(text, level=1)
    return p

def h2(text):
    return doc.add_heading(text, level=2)

def h3(text):
    return doc.add_heading(text, level=3)

def body(text):
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.size = Pt(11)
        run.font.name = "맑은 고딕"
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.35
    return p

def bullet(text):
    p = doc.add_paragraph(text, style="List Bullet")
    for run in p.runs:
        run.font.size = Pt(11)
        run.font.name = "맑은 고딕"
    return p

def code_block(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.font.name = "Consolas"
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(10)
    return p

# ========== COVER ==========
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("\n\n\n실전 따라 하기\n")
r.bold = True
r.font.size = Pt(28)
r.font.name = "맑은 고딕"

t2 = doc.add_paragraph()
t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = t2.add_run("파이썬으로 엑셀 반복 업무 자동화")
r2.bold = True
r2.font.size = Pt(22)
r2.font.name = "맑은 고딕"

t3 = doc.add_paragraph()
t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = t3.add_run("\n시트 병합 · KPI 집계 · 데이터 정제\n예제 3종 포함\n\n")
r3.font.size = Pt(14)

t4 = doc.add_paragraph()
t4.alignment = WD_ALIGN_PARAGRAPH.CENTER
r4 = t4.add_run("ONHAE  ·  전자책 원고 초안\n")
r4.font.size = Pt(12)

t5 = doc.add_paragraph()
t5.alignment = WD_ALIGN_PARAGRAPH.CENTER
r5 = t5.add_run("본 문서는 크몽 등 판매용 전자책 초안입니다.\n예제 파일과 함께 사용하세요.\n")
r5.font.size = Pt(10)

doc.add_page_break()

# ========== TOC-like ==========
h1("목차")
body("1장  이 책을 시작하기 전에")
body("2장  환경 준비 (Python · openpyxl)")
body("3장  실전 CASE 01 — 지점별 시트 자동 병합")
body("4장  실전 CASE 02 — 주문 RAW를 일별 KPI로")
body("5장  실전 CASE 03 — 경비 원장 정제와 보고")
body("6장  자주 하는 실수와 해결")
body("7장  업무에 적용하는 체크리스트")
body("8장  외주로 맡길 때 요청 템플릿")
body("부록  용어 · 다음 단계 · 면책")
doc.add_page_break()

# ========== CH1 ==========
h1("1장  이 책을 시작하기 전에")
h2("1.1 이 책이 다루는 문제")
body("많은 직장과 사장님 업무에서 엑셀은 여전히 중심 도구입니다. 문제는 엑셀 자체가 아니라, 매주 같은 일을 손으로 반복한다는 점입니다. 지점별로 나뉜 시트를 합치고, 주문 원장에서 일별 매출을 뽑고, 지저분한 경비 데이터를 정리하는 일은 한 번이면 간단해 보여도 열 번, 스무 번 반복되면 시간과 실수가 쌓입니다.")
body("이 책은 파이썬 라이브러리 openpyxl을 사용해, 그런 반복 작업을 스크립트로 바꾸는 실전 방법을 다룹니다. 이론 나열보다 따라 하면 결과가 나오는 예제 세 가지를 중심으로 구성했습니다.")

h2("1.2 이 책으로 할 수 있는 것")
bullet("여러 시트에 흩어진 데이터를 하나의 통합 원장으로 합치기")
bullet("헤더 이름과 날짜 형식이 제각각일 때 정규화하기")
bullet("주문 RAW에서 취소 건을 제외한 일별 KPI 만들기")
bullet("금액·상태 표기가 지저분한 원장을 보고용으로 정제하기")
bullet("중복 행을 표시하고 합계에서 제외하기")

h2("1.3 이 책이 다루지 않는 것")
bullet("엑셀 기초 함수 전체 강의 (필요 시 별도 학습)")
bullet("VBA 매크로 문법 전 과정")
bullet("대규모 ERP·웹 시스템 구축")
bullet("불법 수집, 계정 무단 접근, 스팸성 자동화")

h2("1.4 대상 독자")
body("엑셀은 쓰지만 매주 같은 마감에 지친 분, 파이썬을 조금 알거나 처음 설치해 볼 의지가 있는 분, 외주를 주기 전에 범위를 이해하고 싶은 분에게 맞습니다. 프로그래밍 전문가가 아니어도 됩니다. 대신 예제 폴더를 열고, 안내에 따라 실행해 보는 태도가 필요합니다.")

h2("1.5 권장 학습 순서")
body("1장과 2장을 읽은 뒤, CASE 01부터 순서대로 따라 하세요. 각 CASE는 독립적으로도 읽을 수 있지만, 환경 준비는 2장이 공통입니다. 막히면 6장 실수 모음을 먼저 보세요.")
doc.add_page_break()

# ========== CH2 ==========
h1("2장  환경 준비")
h2("2.1 필요한 것")
bullet("Windows PC (이 책 예제는 Windows 기준으로 설명)")
bullet("Microsoft Excel 또는 예제 xlsx를 열 수 있는 프로그램")
bullet("Python 3.10 이상 권장")
bullet("인터넷 (최초 패키지 설치 시)")

h2("2.2 Python 설치 확인")
body("명령 프롬프트 또는 PowerShell을 열고 아래를 입력합니다.")
code_block("python --version")
body("버전이 표시되면 설치되어 있는 것입니다. 안 되면 python.org 에서 Windows 설치 파일을 받아 설치하고, 설치 중 Add Python to PATH 옵션을 켭니다.")

h2("2.3 openpyxl 설치")
code_block("python -m pip install openpyxl")
body("설치가 끝나면 파이썬에서 import openpyxl 이 오류 없이 동작해야 합니다.")

h2("2.4 작업 폴더 만들기")
body("예: 바탕화면에 excel_auto 폴더를 만들고, 그 안에 case01, case02, case03 하위 폴더를 둡니다. 각 폴더에 BEFORE 엑셀과 run_xxx.py 스크립트를 넣어 두고 실행합니다.")
code_block("excel_auto/\n  case01/\n  case02/\n  case03/")

h2("2.5 스크립트 실행 방법")
body("해당 폴더에서 다음처럼 실행합니다.")
code_block("cd 경로\\excel_auto\\case01\npython run_merge.py")
body("성공하면 AFTER 파일이 생성되거나 기존 AFTER를 덮어쓰는 방식으로 결과가 나옵니다. 경로는 본인 PC에 맞게 바꿉니다.")

h2("2.6 안전 수칙")
bullet("중요한 원본은 복사본으로 작업하세요.")
bullet("회사 파일이면 개인정보·비밀정보를 가린 샘플로 먼저 연습하세요.")
bullet("스크립트가 쓰는 파일명·시트명을 예제와 다르게 바꿨다면 코드 안 경로도 함께 수정하세요.")
doc.add_page_break()

# ========== CH3 ==========
h1("3장  실전 CASE 01 — 지점별 시트 자동 병합")
h2("3.1 상황")
body("카페나 매장이 여러 곳이면, 지점마다 매출 시트를 따로 관리하는 경우가 많습니다. 마감 때 복사해서 한 장으로 합치다 보면 헤더 이름이 미묘하게 다르거나, 날짜가 2026-03-01 과 2026/03/01 처럼 섞여 집계가 깨집니다.")

h2("3.2 목표")
bullet("지점별 시트를 하나의 통합원장으로 합친다")
bullet("헤더를 표준 이름으로 맞춘다")
bullet("날짜 형식을 통일한다")
bullet("지점별 요약(총매출, 주문수, 객단가, 비중)을 만든다")

h2("3.3 BEFORE 파일의 특징")
body("예제 BEFORE 파일에는 강남점, 홍대점, 판교점, 수원점 시트가 있습니다. 판교점은 헤더 이름이 다른 시트와 다릅니다. 메모 시트는 병합 대상에서 제외합니다.")

h2("3.4 핵심 아이디어: 헤더 매핑")
body("시트마다 다른 한글 헤더를 내부 표준 키로 바꿉니다. 예를 들어 일자와 날짜는 모두 date 로, 매출액과 매출은 모두 sales 로 취급합니다.")
code_block("HEADER_MAP = {\n  \"날짜\": \"date\", \"일자\": \"date\",\n  \"지점\": \"store\", \"매장명\": \"store\",\n  \"카테고리\": \"category\", \"품목군\": \"category\",\n  \"매출\": \"sales\", \"매출액\": \"sales\",\n  \"건수\": \"count\", \"주문수\": \"count\",\n}")

h2("3.5 처리 파이프라인")
body("1) 메모로 시작하는 시트를 제외하고 시트를 순회한다. 2) 1행 헤더를 읽어 매핑한다. 3) 2행부터 데이터를 읽어 표준 키 딕셔너리로 담는다. 4) 날짜 문자열을 파싱한다. 5) 정렬 후 통합원장 시트에 쓴다. 6) 지점별 합계와 비중을 계산한다.")

h2("3.6 따라 하기")
body("1. case01 폴더의 BEFORE 파일을 엑셀로 열어 구조를 확인합니다. 2. run_merge.py 를 실행합니다. 3. 생성된 AFTER 파일에서 통합원장과 지점별 요약을 확인합니다. 4. 본인 데이터에 적용할 때는 시트 이름 규칙과 헤더 매핑 표만 수정하면 됩니다.")

h2("3.7 응용 아이디어")
bullet("월별 파일 여러 개를 폴더 단위로 읽어 합치기")
bullet("카테고리 요약 시트 추가")
bullet("특정 지점만 필터링한 리포트 시트 만들기")

h2("3.8 점검 체크리스트")
bullet("메모·참고 시트가 합쳐지지 않았는가")
bullet("날짜가 텍스트로 남아 정렬이 이상하지 않은가")
bullet("매출 합계가 원본 지점 합과 대략 일치하는가")
doc.add_page_break()

# ========== CH4 ==========
h1("4장  실전 CASE 02 — 주문 RAW를 일별 KPI로")
h2("4.1 상황")
body("스마트스토어·쿠팡·자사몰 등에서 내려받은 주문 원장(RAW)은 행이 깁니다. 보고에 필요한 숫자는 매일 피벗으로 다시 뽑는 경우가 많고, 취소 주문을 매출에 넣을지 말지 사람마다 기준이 달라지기도 합니다.")

h2("4.2 목표")
bullet("취소 주문을 유효 매출에서 제외하는 규칙을 코드로 고정")
bullet("일별 주문수, 유효주문, 취소수, 매출, 취소율, 객단가 산출")
bullet("채널별 요약, 상품 TOP, 한 장 대시보드 요약")

h2("4.3 비즈니스 규칙을 먼저 문장으로")
body("자동화 전에 규칙을 한 문장으로 적으세요. 예: 상태가 취소인 행은 유효 매출과 유효 주문 수에서 제외한다. 취소율은 전체 주문 수 대비 취소 건수의 비율로 계산한다. 객단가는 유효 매출을 유효 주문 수로 나눈다.")

h2("4.4 처리 파이프라인")
body("1) RAW 시트를 읽는다. 2) 수량과 단가로 금액 필드를 만든다. 3) 상태 값으로 유효/취소를 구분한다. 4) 날짜별로 집계한다. 5) 채널별·상품별로 집계한다. 6) 요약 지표를 대시보드 시트에 적는다.")

h2("4.5 따라 하기")
body("case02 폴더의 BEFORE 주문 RAW를 연 뒤, run_daily_kpi.py 를 실행합니다. AFTER 파일의 일별 KPI 시트에서 날짜별 숫자가 채워졌는지 확인합니다. 본인 RAW는 열 위치가 다를 수 있으므로, 스크립트의 열 번호 또는 헤더 이름을 맞춰야 합니다.")

h2("4.6 응용")
bullet("주간 단위 롤업")
bullet("전주 대비 증감 열 추가")
bullet("특정 채널만 남긴 경영 보고용 시트")
doc.add_page_break()

# ========== CH5 ==========
h1("5장  실전 CASE 03 — 경비 원장 정제와 보고")
h2("5.1 상황")
body("경비 청구 데이터는 사람이 여러 방식으로 입력합니다. 12000, 12,000, 12000원 이 섞이고, 승인·승인완료·뒤에 공백이 있는 대기 가 섞이며, 날짜도 형식 이 제각각입니다. 같은 내용이 두 번 들어가면 합계가 과대해집니다.")

h2("5.2 목표")
bullet("금액 파서로 숫자 통일")
bullet("날짜 파서로 표준 날짜 통일")
bullet("상태 표준화 (승인, 대기, 반려, 미제출)")
bullet("중복 플래그")
bullet("보고 포함 조건: 승인 이면서 중복이 아님")
bullet("부서별·항목별 승인 합계와 품질 검수 시트")

h2("5.3 파서 개념")
body("파서는 지저분한 입력을 약속된 형식으로 바꾸는 작은 규칙 모음입니다. 완벽한 인공지능이 아니라, 우리 업무에서 자주 나오는 패턴을 처리하면 충분합니다.")

h2("5.4 중복 키")
body("예제에서는 날짜, 부서, 이름, 항목, 금액, 상태가 모두 같으면 중복 후보로 봅니다. 업무에 따라 전표번호가 있으면 그 키를 쓰는 편이 더 정확합니다.")

h2("5.5 따라 하기")
body("case03 폴더의 BEFORE 파일을 확인한 뒤 정제 스크립트를 실행합니다. AFTER의 정제원장에서 중복 플래그와 보고포함 열을 보고, 부서별 승인합계가 기대와 맞는지 검토합니다.")

h2("5.6 실무 팁")
bullet("회계 제출 전에 품질 검수 시트를 한 번 더 본다")
bullet("승인 총액을 원본 수동 합과 비교한다")
bullet("규칙은 문서나 주석으로 남겨 담당자가 바뀌어도 유지되게 한다")
doc.add_page_break()

# ========== CH6 ==========
h1("6장  자주 하는 실수와 해결")
h2("6.1 ModuleNotFoundError: openpyxl")
body("pip 설치가 다른 파이썬에 되었을 수 있습니다. python -m pip install openpyxl 로 현재 실행 중인 파이썬에 설치하세요.")

h2("6.2 파일이 없다 / 경로 오류")
body("스크립트와 엑셀이 같은 폴더에 있는지 확인합니다. 한글 경로 문제는 드물지만, 문제가 있으면 영문 폴더로 옮겨 테스트합니다.")

h2("6.3 결과가 비어 있다")
body("헤더 매핑이 안 맞아 열이 무시되었을 수 있습니다. 1행 헤더를 출력해 보거나, 매핑 표를 원본에 맞게 수정합니다.")

h2("6.4 날짜가 이상하다")
body("이미 날짜 타입인 셀과 문자열 셀이 섞여 있을 수 있습니다. 파서에서 두 경우를 모두 처리하는지 확인합니다.")

h2("6.5 숫자에 콤마가 남아 계산이 안 된다")
body("금액 파서에서 콤마와 원 문자를 제거한 뒤 int 로 변환하세요.")

h2("6.6 엑셀에서 파일이 열려 있으면 저장 실패")
body("AFTER 파일을 엑셀에서 닫고 다시 실행합니다.")
doc.add_page_break()

# ========== CH7 ==========
h1("7장  업무에 적용하는 체크리스트")
h2("7.1 자동화 후보를 고르는 기준")
bullet("한 달에 2회 이상 반복되는가")
bullet("실수하면 금액·보고에 영향이 있는가")
bullet("규칙이 문장으로 설명 가능한가")
bullet("입출력 파일이 비교적 고정적인가")

h2("7.2 적용 전")
bullet("원본 백업")
bullet("규칙 3~10줄 문서화")
bullet("샘플로 먼저 검증")

h2("7.3 적용 후")
bullet("합계 대사 (원본 vs 결과)")
bullet("이상치 몇 건 샘플 확인")
bullet("실행 방법을 팀 공유")
doc.add_page_break()

# ========== CH8 ==========
h1("8장  외주로 맡길 때 요청 템플릿")
body("직접 하기 어려우면 전문가에게 맡길 수 있습니다. 아래를 채워 보내면 견적이 빨라집니다.")
body("1) 현재 파일 형태 (시트 수, 대략 행 수)")
body("2) 원하는 결과 (예시 파일이나 설명)")
body("3) 포함·제외 규칙")
body("4) 희망 기한")
body("5) Windows 엑셀 여부")
body("이 책의 예제 구조는 외주 범위를 설명하는 샘플로도 활용할 수 있습니다.")
doc.add_page_break()

# ========== APPENDIX ==========
h1("부록")
h2("A. 용어")
body("RAW: 가공 전 원본 데이터. KPI: 핵심 성과 지표. 정규화: 제각각인 형식을 하나로 맞춤. 파서: 입력을 표준 형식으로 바꾸는 규칙.")
h2("B. 다음 단계")
body("폴더 단위 일괄 처리, 주간 리포트 메일 직전 단계까지 자동화, 간단한 실행 배치 파일 만들기 등을 이어서 연습할 수 있습니다.")
h2("C. 면책")
body("본 문서와 예제는 교육·실습 목적의 가상 데이터를 사용합니다. 회사 데이터에 적용할 때는 내부 규정과 개인정보 보호를 지키세요. 스크립트 실행 결과는 사용자 환경에 따라 달라질 수 있으며, 중요 업무 적용 전 충분한 검증이 필요합니다. 판매용으로 재편집할 때는 최신 예제 파일과 함께 제공하세요.")
h2("D. 판권·버전")
body("초안 버전 0.1  ·  ONHAE")
body("제목: 파이썬으로 엑셀 반복 업무 자동화 실전")
body("구성: 실전 CASE 3종 중심 전자책 원고")

doc.save(docx_path)
print("DOCX saved", docx_path, "exists", docx_path.exists())

# Convert via HWP COM
pythoncom.CoInitialize()
hwp = None
try:
    hwp = win32com.client.Dispatch("HWPFrame.HwpObject")
    # silence security popup if module registered - try common names
    try:
        hwp.RegisterModule("FilePathCheckDLL", "FilePathCheckerModule")
    except Exception:
        try:
            hwp.RegisterModule("FilePathCheckDLL", "FilePathCheckerModuleExample")
        except Exception as e:
            print("RegisterModule skip", e)
    try:
        hwp.XHwpWindows.Item(0).Visible = False
    except Exception:
        pass

    # Try open docx
    opened = False
    for fmt_try in [str(docx_path), str(docx_path.resolve())]:
        try:
            # Open method signatures vary
            ret = hwp.Open(fmt_try)
            print("Open ret", ret, fmt_try)
            opened = True
            break
        except Exception as e:
            print("Open fail", e)

    if not opened:
        # create new and paste text from plain
        hwp.HAction.Run("FileNew")
        # read docx paragraphs roughly via plain export
        from docx import Document as D2
        d2 = D2(str(docx_path))
        full = []
        for p in d2.paragraphs:
            full.append(p.text)
        text = "\n".join(full)
        # insert in chunks
        chunk = 3000
        for i in range(0, len(text), chunk):
            part = text[i:i+chunk]
            hwp.HAction.GetDefault("InsertText", hwp.HParameterSet.HInsertText.HSet)
            hwp.HParameterSet.HInsertText.Text = part
            hwp.HAction.Execute("InsertText", hwp.HParameterSet.HInsertText.HSet)
        print("Inserted text length", len(text))

    # Save as HWP
    save_path = str(hwp_path.resolve())
    # Clear if exists
    if hwp_path.exists():
        hwp_path.unlink()
    try:
        hwp.SaveAs(save_path, "HWP")
        print("SaveAs HWP ok")
    except Exception as e:
        print("SaveAs HWP fail", e)
        try:
            hwp.SaveAs(save_path)
            print("SaveAs default ok")
        except Exception as e2:
            print("SaveAs fail2", e2)
            # try format code
            try:
                hwp.HAction.GetDefault("FileSaveAs_S", hwp.HParameterSet.HFileOpenSave.HSet)
                hwp.HParameterSet.HFileOpenSave.filename = save_path
                hwp.HParameterSet.HFileOpenSave.Format = "HWP"
                hwp.HAction.Execute("FileSaveAs_S", hwp.HParameterSet.HFileOpenSave.HSet)
                print("FileSaveAs_S tried")
            except Exception as e3:
                print("FileSaveAs_S fail", e3)

    time.sleep(1)
    print("HWP exists", hwp_path.exists(), hwp_path)
    if hwp_path.exists():
        print("HWP size", hwp_path.stat().st_size)
finally:
    if hwp is not None:
        try:
            hwp.Quit()
        except Exception:
            pass
    pythoncom.CoUninitialize()
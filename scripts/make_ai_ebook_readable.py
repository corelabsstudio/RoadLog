# -*- coding: utf-8 -*-
"""
AI 부업 준비 전자책 — 여백 최소화 + 가독성 강화 판형
- 쪽 안 빈 공간 축소 (크몽 반려 대응)
- 제목/부제/본문 계층 명확
- DOCX + HWP 출력
"""
from __future__ import annotations

import time
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor

ROOT = Path.home() / "Desktop" / "크몽_전자책"
DIR_MS = ROOT / "01_원고"
DIR_MS.mkdir(parents=True, exist_ok=True)

DOCX = DIR_MS / "AI로_부업_준비하기_가독성판.docx"
HWP = DIR_MS / "AI로_부업_준비하기_가독성판.hwp"

# colors
C_TITLE = RGBColor(15, 23, 42)
C_H1 = RGBColor(30, 64, 175)
C_H2 = RGBColor(30, 64, 175)
C_H3 = RGBColor(51, 65, 85)
C_BODY = RGBColor(30, 41, 59)
C_MUTED = RGBColor(71, 85, 105)
C_ACCENT = RGBColor(37, 99, 235)


def set_run_font(run, name="맑은 고딕", size=10.5, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    rFonts.set(qn("w:eastAsia"), name)
    rFonts.set(qn("w:cs"), name)


def set_para_dense(p, after=4, before=0, line=1.12, align=None):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    try:
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    except Exception:
        pass
    if align is not None:
        p.alignment = align


def shade_paragraph(paragraph, fill_hex="EEF2FF"):
    """Light background on paragraph for subtitle bars."""
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill_hex)
    shd.set(qn("w:val"), "clear")
    pPr.append(shd)


def add_bottom_border(paragraph, color="2563EB", sz="12"):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), sz)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def build() -> Path:
    doc = Document()
    # tight margins — less empty edge
    for sec in doc.sections:
        sec.top_margin = Cm(1.5)
        sec.bottom_margin = Cm(1.5)
        sec.left_margin = Cm(1.8)
        sec.right_margin = Cm(1.8)
        sec.page_width = Cm(21.0)
        sec.page_height = Cm(29.7)

    def title_block(main: str, sub: str = ""):
        p = doc.add_paragraph()
        set_para_dense(p, after=2, before=0, line=1.1)
        r = p.add_run(main)
        set_run_font(r, size=22, bold=True, color=C_TITLE)
        if sub:
            p2 = doc.add_paragraph()
            set_para_dense(p2, after=6, before=2, line=1.1)
            r2 = p2.add_run(sub)
            set_run_font(r2, size=12, bold=False, color=C_MUTED)
            shade_paragraph(p2, "EEF2FF")

    def h1(text: str, subtitle: str = ""):
        # small spacer only
        p = doc.add_paragraph()
        set_para_dense(p, after=2, before=8, line=1.05)
        r = p.add_run(text)
        set_run_font(r, size=15, bold=True, color=C_H1)
        add_bottom_border(p, "3B82F6", "18")
        if subtitle:
            p2 = doc.add_paragraph()
            set_para_dense(p2, after=6, before=3, line=1.1)
            r2 = p2.add_run(subtitle)
            set_run_font(r2, size=10, bold=True, color=C_ACCENT)
            shade_paragraph(p2, "F1F5F9")

    def h2(text: str):
        p = doc.add_paragraph()
        set_para_dense(p, after=3, before=7, line=1.1)
        r = p.add_run("■ " + text)
        set_run_font(r, size=12, bold=True, color=C_H2)

    def h3(text: str):
        p = doc.add_paragraph()
        set_para_dense(p, after=2, before=5, line=1.1)
        r = p.add_run("▸ " + text)
        set_run_font(r, size=11, bold=True, color=C_H3)

    def body(text: str):
        p = doc.add_paragraph()
        set_para_dense(p, after=3, before=0, line=1.15)
        r = p.add_run(text)
        set_run_font(r, size=10.5, color=C_BODY)
        pf = p.paragraph_format
        pf.first_line_indent = Cm(0)  # no indent waste on short paras; use block
        return p

    def bullet(text: str):
        p = doc.add_paragraph()
        set_para_dense(p, after=2, before=0, line=1.12)
        r = p.add_run("· " + text)
        set_run_font(r, size=10.5, color=C_BODY)
        p.paragraph_format.left_indent = Cm(0.3)

    def num_item(n: int, text: str):
        p = doc.add_paragraph()
        set_para_dense(p, after=2, before=0, line=1.12)
        r = p.add_run(f"{n}. {text}")
        set_run_font(r, size=10.5, color=C_BODY)
        p.paragraph_format.left_indent = Cm(0.25)

    def callout(title: str, text: str, fill="FEF3C7"):
        p = doc.add_paragraph()
        set_para_dense(p, after=1, before=4, line=1.1)
        r = p.add_run("【" + title + "】")
        set_run_font(r, size=10, bold=True, color=C_TITLE)
        shade_paragraph(p, fill)
        p2 = doc.add_paragraph()
        set_para_dense(p2, after=5, before=0, line=1.12)
        r2 = p2.add_run(text)
        set_run_font(r2, size=10, color=C_BODY)
        shade_paragraph(p2, fill)

    def prompt_box(text: str):
        p = doc.add_paragraph()
        set_para_dense(p, after=4, before=2, line=1.12)
        r = p.add_run(text)
        set_run_font(r, name="Consolas", size=9, color=C_BODY)
        shade_paragraph(p, "F8FAFC")

    def kv(label: str, text: str):
        p = doc.add_paragraph()
        set_para_dense(p, after=2, before=0, line=1.12)
        r1 = p.add_run(label + "  ")
        set_run_font(r1, size=10.5, bold=True, color=C_ACCENT)
        r2 = p.add_run(text)
        set_run_font(r2, size=10.5, color=C_BODY)

    # ========== COVER (compact, no huge blank) ==========
    p = doc.add_paragraph()
    set_para_dense(p, after=4, before=2, line=1.05, align=WD_ALIGN_PARAGRAPH.CENTER)
    r = p.add_run("ONHAE 전자책 초안  v0.2 가독성·밀도 강화판")
    set_run_font(r, size=9, color=C_MUTED)

    p = doc.add_paragraph()
    set_para_dense(p, after=2, before=6, line=1.05, align=WD_ALIGN_PARAGRAPH.CENTER)
    r = p.add_run("개발 몰라도 따라 하는")
    set_run_font(r, size=13, bold=True, color=C_MUTED)

    p = doc.add_paragraph()
    set_para_dense(p, after=4, before=0, line=1.05, align=WD_ALIGN_PARAGRAPH.CENTER)
    r = p.add_run("AI로 부업 준비하기")
    set_run_font(r, size=24, bold=True, color=C_TITLE)

    p = doc.add_paragraph()
    set_para_dense(p, after=6, before=2, line=1.1, align=WD_ALIGN_PARAGRAPH.CENTER)
    r = p.add_run("실전 노트  ·  템플릿  ·  체크리스트")
    set_run_font(r, size=12, bold=True, color=C_ACCENT)
    shade_paragraph(p, "EEF2FF")

    callout(
        "이 책의 성격",
        "수익 자랑기가 아닙니다. AI로 부업 상품·판매 글을 준비하는 과정을 쉬운 말로 정리한 실습 노트입니다. 수익은 보장되지 않습니다. 실행량이 결과를 만듭니다.",
        "EEF2FF",
    )
    body("대상: 개발은 모르지만 마우스·키보드·회원가입·파일 저장은 가능한 분. 말투: 친절하고 짧게. 목표: 한 주 안에 상품 문장과 등록 준비를 끝낼 수 있는 도구 제공.")
    kv("포함", "본문 전 장 + 별도 폴더의 템플릿·체크리스트(02_부록, 03_체크리스트)")
    kv("주의", "허위 수익·불법 자동화·스팸 안내는 다루지 않습니다.")

    # ========== TOC dense ==========
    h1("목차", "순서대로 읽어도 되고, 필요한 장만 펼쳐도 됩니다")
    toc = [
        ("머리말", "왜 이 책을 썼는지"),
        ("1장", "약속과 한계 — 무엇을 주고 무엇을 안 주는지"),
        ("2장", "AI를 아주 쉽게 이해하기"),
        ("3장", "개발 모르는 분 기준의 쉬움"),
        ("4장", "부업 아이디어를 AI로 고르기"),
        ("5장", "팔 수 있는 형태로 다듬기"),
        ("6장", "판매 글·FAQ·제목을 AI로 쓰기"),
        ("7장", "포트폴리오와 이미지"),
        ("8장", "크몽·숨고 흐름"),
        ("9장", "하면 안 되는 것"),
        ("10장", "일주일 실행 계획"),
        ("11장", "준비 과정 사례 노트"),
        ("12장", "자주 묻는 질문"),
        ("맺음말", "오늘 할 일 세 가지"),
    ]
    for a, b in toc:
        p = doc.add_paragraph()
        set_para_dense(p, after=2, before=0, line=1.1)
        r1 = p.add_run(f"{a}  ")
        set_run_font(r1, size=10.5, bold=True, color=C_H1)
        r2 = p.add_run(b)
        set_run_font(r2, size=10.5, color=C_BODY)

    # ========== 머리말 ==========
    h1("머리말", "큰돈 인증이 아니라, 준비 속도를 높이는 책")
    body("AI 이야기는 어디에나 있습니다. 유튜브와 전자책에는 금방 돈을 벌 수 있다는 말도 많습니다. 그런데 막상 시작하려면 어디서부터 눌러야 할지 막막합니다.")
    body("저는 개발자가 아니어도 할 수 있는 일부터 AI에게 물어 가며 준비했습니다. 서비스 설명 쓰기, 자주 묻는 질문 만들기, 가격 나누기, 포트폴리오 구성, 프로필 사진과 한 줄 소개 정하기 같은 일입니다.")
    body("느낀 점은 분명합니다. AI는 마법이 아닙니다. 대신 빈 화면에서 혼자 고민하는 시간을 크게 줄여 줍니다. 이 책은 그 방법을 어려운 용어 없이 순서대로 적었습니다.")
    body("아직 큰 수익을 인증할 단계는 아닙니다. 그래서 수익 자랑 대신, 바로 써먹을 문장과 체크리스트를 드립니다. 같은 출발선에 있는 분에게 도움이 되기를 바랍니다.")
    callout("한 줄 요약", "AI = 자동 입금 기계가 아니라, 부업 준비 초안을 빠르게 만들어 주는 도우미.", "ECFDF5")

    # ========== 1 ==========
    h1("1장  이 책의 약속과 한계", "기대치를 맞춰야 후회가 없습니다")
    h2("1.1 약속하는 것")
    bullet("AI 채팅을 부업 준비에 쓰는 구체적인 순서")
    bullet("복사해서 고칠 수 있는 문장 템플릿(부록 폴더)")
    bullet("한 주 동안 따라 할 수 있는 할 일 표")
    bullet("과장 광고에 속지 않기 위한 기준")
    h2("1.2 약속하지 않는 것")
    bullet("특정 기간 안에 무조건 수익")
    bullet("하루 만에 부자 되는 비법")
    bullet("불법·편법·스팸으로 돈 버는 방법")
    bullet("코딩 전문가가 되는 완전한 강의")
    h2("1.3 추천 독자")
    body("마우스와 키보드로 파일을 저장하고, 인터넷 검색과 회원가입 정도는 할 수 있는 분. 스마트폰이나 PC로 AI 화면에 글을 입력해 본 적이 있거나, 앞으로 해 볼 의지가 있는 분.")
    body("전원을 켜는 것부터 어렵다면, 먼저 파일·폴더·회원가입만 익힌 뒤 이 책을 보세요. 이 책은 ‘개발 몰라도 되는 쉬운 단계’를 목표로 하되, 기본 조작은 가능하다고 가정합니다.")
    h2("1.4 읽기 전 준비물")
    num_item(1, "인터넷이 되는 컴퓨터 또는 노트북")
    num_item(2, "이메일 계정")
    num_item(3, "AI 채팅 도구 하나(무료부터 가능)")
    num_item(4, "메모장·한글·워드 중 하나")
    num_item(5, "메모할 노트 앱 또는 종이")
    callout("쉬움의 기준", "클릭·타이핑·저장·회원가입·결제 창 열기 가능. 프로그래밍 언어 암기 불필요.", "F1F5F9")

    # ========== 2 ==========
    h1("2장  AI를 아주 쉽게 이해하기", "검색과 비슷하지만, 글을 이어서 다듬을 수 있다")
    h2("2.1 한 문장 정의")
    body("AI 채팅은 내가 질문이나 글을 넣으면, 그에 맞는 초안·아이디어·정리본을 만들어 주는 도우미입니다. 검색처럼 정보를 찾기도 하고, 여러 문장을 원하는 말투로 다시 써 달라고 부탁할 수도 있습니다.")
    h2("2.2 잘 하는 일")
    bullet("긴 설명을 짧게 줄이기")
    bullet("판매 글·FAQ 초안 쓰기")
    bullet("단계·표·목록으로 나누기")
    bullet("어려운 말을 쉬운 말로 바꾸기")
    bullet("빠진 항목을 찾아 달라고 하기")
    h2("2.3 그대로 믿으면 위험한 일")
    bullet("법률·세금·의료를 최종 결정으로 믿기")
    bullet("확인 없이 숫자·사례를 사실로 쓰기")
    bullet("비밀번호·주민번호 등 민감 정보를 붙여 넣기")
    bullet("검토 없이 코드를 실행하기")
    h2("2.4 기본 사용 5단계")
    num_item(1, "하고 싶은 일을 한 줄로 쓴다. 예: 엑셀 자동화 서비스 소개 문단 써 줘.")
    num_item(2, "누구를 위한 글인지 적는다. 예: 개발 모르는 직장인.")
    num_item(3, "길이와 말투를 정한다. 예: 200자, 친절한 말투, 과한 광고 말투 금지.")
    num_item(4, "나온 글을 읽고 사실과 다른 부분을 고친다.")
    num_item(5, "다시 부탁한다. 예: 더 쉽게, 목록으로, 30자 이내 제목으로.")
    h2("2.5 복붙용 기본 질문 틀")
    prompt_box(
        "역할: 너는 초보도 이해하게 설명하는 도우미다.\n"
        "목표: (하고 싶은 일)\n"
        "독자: (누구를 위한 글인지)\n"
        "조건: 쉬운 말, 과장 금지, 단계로 나눠 줘.\n"
        "출력: (문단 / 목록 / 표 중 선택)"
    )
    body("위 틀은 부록 파일 02_AI_질문_기본틀.txt 와 같습니다. 폴더에서 복사해 쓰세요.")

    # ========== 3 ==========
    h1("3장  개발 모르는 분 기준", "막히면 멈추는 기준도 함께")
    h2("3.1 이 책의 쉬움")
    body("클릭, 타이핑, 파일 저장, 회원가입, 결제 창 열기 정도는 가능하다고 가정합니다. 프로그래밍 문법을 외울 필요는 없습니다.")
    h2("3.2 어려우면 여기서 멈추기")
    bullet("안내와 전혀 다른 화면이 나와 30분 이상 혼자 헤맬 때 → 캡처 저장 후 다음 날 재시도하거나 주변에 묻기")
    bullet("결제·개인정보 창이 의심스러울 때 → 진행하지 않기")
    h2("3.3 폴더 하나 만들기")
    body("바탕화면에 ‘부업준비’ 폴더를 만드세요. 그 안에 글초안, 이미지, 체크리스트 하위 폴더를 두면 파일이 섞이지 않습니다. 이 전자책 폴더(크몽_전자책)와 옆에 나란히 두어도 좋습니다.")

    # ========== 4 ==========
    h1("4장  부업 아이디어를 AI로 고르기", "많을수록 좋은 게 아니라, 줄이는 게 실력")
    h2("4.1 세 가지 필터")
    num_item(1, "내가 이미 조금 해 본 일, 또는 배울 의지가 있는 일인가")
    num_item(2, "다른 사람이 돈·시간을 아끼고 싶어 하는 일인가")
    num_item(3, "불법·사기·스팸에 가깝지 않은가")
    body("세 가지를 통과한 것만 남기세요. 10개를 받는 이유는 고르기 위해서입니다.")
    h2("4.2 AI에게 물을 질문 예시")
    prompt_box(
        "나는 Windows를 쓰고 엑셀을 조금 다룰 수 있다. 개발은 모른다. "
        "원격으로 할 수 있는 부업 아이디어 10개를 난이도 하·중으로만 짧게 나열해 줘. "
        "각각 한 줄로 왜 돈이 될 수 있는지 적어 줘. 과장 없이."
    )
    h2("4.3 3개로 줄인 뒤")
    prompt_box(
        "이 아이디어로 첫 상품을 만든다면, 팔 단위(한 건)를 무엇으로 정하면 좋을까? "
        "초보 기준으로 범위를 좁혀 줘. 아이디어: (선택한 것)"
    )
    h2("4.4 첫 상품은 작게")
    body("처음부터 앱·거대한 사이트를 만들겠다는 목표는 대부분 너무 큽니다. 첫 상품 예: 설명 글 패키지, 템플릿 한 종, 엑셀 정리 한 가지, 체크리스트형 전자책. ‘끝낼 수 있는 단위’가 좋습니다.")
    callout("연습", "지금 빈칸을 채우세요. 나는 (누구) 에게 (결과) 를 (방법) 으로 제공한다.", "FEF3C7")

    # ========== 5 ==========
    h1("5장  팔 수 있는 형태로 다듬기", "한 줄 정의 → 포함/미포함 → 가격 3단")
    h2("5.1 한 줄 정의")
    body("나는 (누구) 에게 (결과) 를 (방법) 으로 제공한다. 예: 나는 엑셀 마감이 반복되는 사람에게, 시트를 합친 결과물을, 원격 작업으로 제공한다. 부록 01_한줄_정의_템플릿.txt 참고.")
    h2("5.2 포함 / 미포함")
    body("나중에 수정 분쟁을 줄입니다. AI 요청 예시는 아래와 같습니다.")
    prompt_box(
        "다음 상품의 포함 항목 5개와 미포함 항목 5개를 적어 줘. "
        "초보 판매자 기준, 과대 약속 금지. 상품: (한 줄 정의)"
    )
    h2("5.3 가격 3단계")
    body("기본(한 가지) · 표준(두 가지 연계) · 고급(복잡 맞춤)으로 나누면 구매자가 비교하기 쉽습니다. 같은 플랫폼에서 비슷한 상품 가격을 여러 개 본 뒤, 덤핑과 과고가를 피하고 가운데를 고르세요.")
    body("처음 가격을 너무 낮추면 주문이 와도 수정 부담에 손해가 날 수 있습니다. 시간은 비용입니다.")
    kv("팁", "기본안 예: 5만 / 10만 / 18만 대 — 시장·본인 일정에 맞게 조정")

    # ========== 6 ==========
    h1("6장  판매 글·FAQ·제목을 AI로 쓰기", "초안은 AI, 최종 결정은 사람")
    h2("6.1 서비스 설명 요청문")
    prompt_box(
        "다음 조건으로 서비스 상세 설명을 써 줘. 200자 이상. 쉬운 말. 과장 금지. "
        "작은따옴표와 꺾쇠는 쓰지 마. 상품: (한 줄 정의). 포함: (목록). "
        "진행 순서: 문의-범위확인-작업-납품."
    )
    h2("6.2 FAQ 요청문")
    prompt_box(
        "초보 구매자가 물을 법한 질문 10개와 짧은 답변을 만들어 줘. "
        "답변 3~5문장 이내. 가격·기간·수정 범위·준비물 포함. 상품: (한 줄 정의)"
    )
    h2("6.3 제목 글자 수 제한")
    body("플랫폼마다 제목 길이 제한이 있습니다. 예: 30자 이내. 제한을 반드시 적어 주세요.")
    prompt_box(
        "다음 의미를 유지한 채 제목 후보 8개를 30자 이내로 만들어 줘. "
        "의미: 엑셀 업무 자동화 맞춤 제작"
    )
    h2("6.4 사람이 반드시 고칠 부분")
    bullet("내가 실제로 못 하는 일 삭제")
    bullet("기간·수정 횟수를 내 일정에 맞게 수정")
    bullet("가짜 경력·가짜 후기 넣지 않기")
    callout("품질 기준", "AI 초안 점수 60점. 내가 고친 뒤 100점. 올리는 건 100점짜리만.", "EEF2FF")

    # ========== 7 ==========
    h1("7장  포트폴리오와 이미지", "수익 인증이 없어도 샘플은 만들 수 있다")
    h2("7.1 왜 필요한가")
    body("처음 보는 사람은 당신을 모릅니다. 샘플 결과물이 있으면 ‘이 사람이 뭘 하는지’가 보입니다. 매출 스크린샷이 없어도, 작업 전후 샘플은 준비할 수 있습니다.")
    h2("7.2 샘플 만들기")
    num_item(1, "가상 데이터로 예제 표 만들기(회사 비밀 데이터 금지)")
    num_item(2, "작업 전·후를 글 또는 이미지로 설명")
    num_item(3, "AI에게 포트폴리오 제목·한 줄 요약 요청")
    num_item(4, "이미지 규격은 플랫폼 안내를 그대로 따르기")
    h2("7.3 업로드가 안 될 때")
    body("규격·용량·브라우저·네트워크 문제일 수 있습니다. 다른 브라우저, 시크릿 창, 아주 단순한 테스트 이미지로 먼저 확인하세요. 임시 저장 후 다시 시도해도 됩니다. 규격 문제와 전송 실패를 구분하면 시간을 아낍니다.")

    # ========== 8 ==========
    h1("8장  크몽·숨고 흐름", "채널마다 글 길이·비용이 다르다")
    h2("8.1 차이 (쉽게)")
    body("크몽은 상품을 미리 올려 두고 구매자가 찾아오는 쪽에 가깝습니다. 숨고는 손님이 요청을 올리면 고수가 견적을 보내는 방식에 가깝습니다. 숨고 견적은 캐시가 필요할 수 있으니, 될 만한 요청만 고르세요.")
    h2("8.2 등록 전 체크")
    bullet("증명할 수 없는 경력·학력 문구 삭제")
    bullet("포함·미포함 명시")
    bullet("수정 규정 작성")
    bullet("작업·연락은 플랫폼 안에서")
    h2("8.3 승인 전후")
    body("승인 전에는 보통 일반 구매 노출이 어렵습니다. 승인 알림을 기다리고, 반려 사유가 오면 그 부분만 고칩니다.")
    h2("8.4 첫 문의가 오면")
    num_item(1, "원하는 결과를 한 줄로 다시 확인")
    num_item(2, "파일·기한·수정 범위 확인")
    num_item(3, "가능하면 작은 범위로 먼저 끝내기")
    num_item(4, "추가 요청은 추가 견적으로 안내")
    body("상세 체크는 03_체크리스트 폴더의 첫문의_응대_체크리스트.txt 를 사용하세요.")

    # ========== 9 ==========
    h1("9장  하면 안 되는 것", "짧게 벌려다 오래 잃는 길")
    bullet("허위 수익·허위 후기")
    bullet("무단 도배·스팸성 홍보 안내")
    bullet("계정·인증 우회, 불법 프로그램 안내")
    bullet("과도한 개인정보를 AI에 붙여 넣기")
    bullet("할 수 없는 일을 할 수 있다고 쓰기")
    body("부업은 오래 가는 쪽이 이득입니다. 계정 정지와 신뢰 손실은 돈으로 잘 안 돌아옵니다.")

    # ========== 10 ==========
    h1("10장  일주일 실행 계획", "완벽한 한 주보다, 멈춘 곳 다시 시작")
    h2("Day 1 — 아이디어")
    body("아이디어 10개 받기 → 3개로 줄이기 → 한 줄 정의 쓰기. 산출물: 메모 한 장.")
    h2("Day 2 — 상품 뼈대")
    body("포함·미포함 각 5줄, 가격 3단계 초안. 산출물: 표 한 장.")
    h2("Day 3 — 판매 글")
    body("상세 설명·FAQ AI 초안 후 직접 수정. 산출물: 등록용 본문.")
    h2("Day 4 — 샘플")
    body("포트폴리오 샘플 1~2개(가상 데이터). 산출물: 이미지 또는 문서.")
    h2("Day 5 — 프로필")
    body("활동명·한 줄 소개·프로필 사진. 산출물: 프로필 완성.")
    h2("Day 6 — 등록")
    body("플랫폼 등록 또는 임시 저장. 막히는 부분 메모.")
    h2("Day 7 — 정리")
    body("전체 글 재독, 과장 삭제, 다음 주 할 일 3개만 정하기.")
    callout("밀렸을 때", "일정을 늘리면 됩니다. 포기만 하지 마세요. 체크리스트 파일로 진행 상황을 표시하세요.", "ECFDF5")

    # ========== 11 ==========
    h1("11장  준비 과정 사례 노트", "매출 자랑이 아니라, 판단 기준 모음")
    body("아래는 큰돈 인증 사례가 아닙니다. 준비하며 마주친 일과 판단 기준입니다. 독자도 자기 일지로 바꿔 쓰세요.")
    h2("11.1 사례 A — 상품 문장")
    kv("상황", "한 페이지 상품 설명이 필요했다.")
    kv("AI", "초안 → 글자 수 제한 재요청 → 못 하는 약속 삭제.")
    kv("결과", "제목·상세·FAQ·절차·준비물 문장 확보.")
    kv("교훈", "첫 초안 그대로 올리지 말 것.")
    h2("11.2 사례 B — 가격")
    kv("상황", "기본·표준·고급 가격이 필요했다.")
    kv("방법", "유사 상품 여러 개 비교, 덤핑·과고가 회피.")
    kv("교훈", "싼 주문이 이득이 아닐 수 있다.")
    h2("11.3 사례 C — 포트폴리오")
    kv("상황", "실거래 전이라 고객 결과물이 부족했다.")
    kv("방법", "가상 데이터 전후 샘플 + 단계 설명.")
    kv("교훈", "수익 인증 없이도 샘플은 가능. 비밀 데이터 금지.")
    h2("11.4 사례 D — 업로드 장애")
    kv("상황", "규격에 맞아 보여도 업로드 실패.")
    kv("대응", "임시 저장, 다른 파일·브라우저, 원인 메모.")
    kv("교훈", "내용 문제와 전송 문제를 분리할 것.")
    h2("11.5 사례 E — 채널 두 개")
    kv("상황", "상품형 플랫폼과 견적형 플랫폼을 함께 씀.")
    kv("대응", "통째 복붙 금지, 견적은 짧은 맞춤 문구, 캐시 드는 쪽은 선별.")
    kv("교훈", "채널마다 문장 길이와 비용 구조가 다르다.")

    # ========== 12 ==========
    h1("12장  자주 묻는 질문", "짧은 답으로 기대치 정리")
    h3("Q1. AI만 있으면 돈이 되나요?")
    body("A. 아닙니다. AI만으로 자동 입금되지 않습니다. 상품·신뢰·실행이 필요합니다. AI는 준비 속도를 높여 줍니다.")
    h3("Q2. 개발을 몰라도 되나요?")
    body("A. 이 책 범위의 글 쓰기·등록·템플릿 활용은 가능합니다. 프로그램을 대신 개발하는 일까지는 배우거나 외주·협업이 필요합니다.")
    h3("Q3. 무료 AI로 충분한가요?")
    body("A. 시작은 무료로 가능합니다. 제한이 느껴지면 나중에 유료를 검토하세요.")
    h3("Q4. 얼마나 걸리나요?")
    body("A. 하루 1~2시간 기준 1~2주면 기본 등록 준비까지 가 보는 경우가 많습니다. 개인 차이가 큽니다.")
    h3("Q5. 수익이 없으면 이 책이 소용 없나요?")
    body("A. 이 책은 수익 대체가 아니라 준비 도구입니다. 실행하지 않으면 어떤 책도 소용 없습니다.")
    h3("Q6. 컴맹도 정말 가능한가요?")
    body("A. ‘개발 몰라도 되는 쉬운 단계’를 목표로 했습니다. 다만 전원·폴더·회원가입이 전혀 안 되면 그 기초를 먼저 익히세요. 이 책은 그 다음 단계입니다.")

    # ========== 맺음말 ==========
    h1("맺음말", "거창한 한 방보다, 작은 상품을 끝까지")
    body("부업은 거창한 한 방이 아니라, 작은 상품을 끝까지 올리는 일의 반복에 가깝습니다. AI는 그 반복에서 글을 다듬고 빈칸을 채워 주는 도구입니다.")
    body("오늘 할 일이 막막하면 세 가지만 하세요. ① 한 줄 정의 쓰기 ② 포함·미포함 다섯 줄 쓰기 ③ AI로 소개 문단 하나 만들기. 그것만으로도 어제보다 한 걸음입니다.")
    body("어색한 문장, 더 쉬운 표현, 추가하고 싶은 사례가 있으면 표시해 두세요. 다음 수정본에 반영할 수 있습니다.")
    p = doc.add_paragraph()
    set_para_dense(p, after=2, before=6, line=1.1)
    r = p.add_run("감사합니다.  ·  ONHAE  ·  AI로 부업 준비하기  ·  v0.2 가독성·밀도 강화판")
    set_run_font(r, size=10, bold=True, color=C_MUTED)

    # fill remaining density: quick reference appendix without huge blank
    h1("빠른 참조 카드", "필요할 때 이 페이지만 펼치세요")
    h2("오늘 할 일 3")
    num_item(1, "한 줄 정의")
    num_item(2, "포함 5 · 미포함 5")
    num_item(3, "소개 문단 1개(AI 초안+수정)")
    h2("등록 전 5")
    bullet("허위 경력 삭제")
    bullet("포함/미포함")
    bullet("수정 규정")
    bullet("제목 글자 수")
    bullet("과장 표현 삭제")
    h2("부록 파일 위치")
    body("02_부록_템플릿 — 문장 틀  |  03_체크리스트 — 실행 표  |  04_표지_안내 — 판매 문구")
    body("끝.")

    doc.save(DOCX)
    print("DOCX", DOCX, DOCX.stat().st_size)
    return DOCX


def convert_hwp(docx_path: Path, hwp_path: Path) -> bool:
    """Insert structured text into HWP with basic char/para spacing via repeated inserts."""
    import pythoncom
    import win32com.client
    from docx import Document as D2

    pythoncom.CoInitialize()
    hwp = None
    try:
        hwp = win32com.client.Dispatch("HWPFrame.HwpObject")
        for mod in ("FilePathCheckerModule", "FilePathCheckerModuleExample"):
            try:
                hwp.RegisterModule("FilePathCheckDLL", mod)
                break
            except Exception:
                pass
        try:
            hwp.XHwpWindows.Item(0).Visible = False
        except Exception:
            pass

        hwp.HAction.Run("FileNew")

        # denser default para: reduce empty feel
        try:
            hwp.HAction.GetDefault("ParagraphShape", hwp.HParameterSet.HParaShape.HSet)
            # LineSpacing might be 160 = 160% ; try 120
            try:
                hwp.HParameterSet.HParaShape.LineSpacing = 120
            except Exception:
                pass
            try:
                hwp.HParameterSet.HParaShape.LineSpacingType = 0  # ratio
            except Exception:
                pass
            hwp.HAction.Execute("ParagraphShape", hwp.HParameterSet.HParaShape.HSet)
        except Exception as e:
            print("parashape", e)

        d2 = D2(str(docx_path))
        parts: list[str] = []
        buf: list[str] = []
        for p in d2.paragraphs:
            t = p.text
            if not t.strip():
                continue  # skip empty paragraphs to reduce blank space in HWP
            buf.append(t)
            if sum(len(x) + 1 for x in buf) > 2000:
                parts.append("\n".join(buf))
                buf = []
        if buf:
            parts.append("\n".join(buf))

        for part in parts:
            hwp.HAction.GetDefault("InsertText", hwp.HParameterSet.HInsertText.HSet)
            hwp.HParameterSet.HInsertText.Text = part + "\n"
            hwp.HAction.Execute("InsertText", hwp.HParameterSet.HInsertText.HSet)

        if hwp_path.exists():
            hwp_path.unlink()
        save_path = str(hwp_path.resolve())
        hwp.HAction.GetDefault("FileSaveAs_S", hwp.HParameterSet.HFileOpenSave.HSet)
        hwp.HParameterSet.HFileOpenSave.filename = save_path
        hwp.HParameterSet.HFileOpenSave.Format = "HWP"
        hwp.HAction.Execute("FileSaveAs_S", hwp.HParameterSet.HFileOpenSave.HSet)
        time.sleep(0.8)
        ok = hwp_path.exists()
        print("HWP", ok, hwp_path.stat().st_size if ok else 0)
        return ok
    finally:
        if hwp is not None:
            try:
                hwp.Quit()
            except Exception:
                pass
        pythoncom.CoUninitialize()


def write_note():
    (DIR_MS / "가독성판_안내.txt").write_text(
        """v0.2 가독성·밀도 강화판 안내
========================
변경 요약
1) 여백: 페이지 여백 축소, 문단 간격·줄간격 축소 → 쪽 안 빈 공간 감소
2) 제목: 장 제목 15pt 굵게 + 아래 파란 선
3) 부제: 장마다 회색/파란 배경 한 줄 부제
4) 본문: 10.5pt, 줄간격 약 1.15, 문단 간격 타이트
5) 소제목: ■ / ▸ 기호로 계층 구분
6) 강조 박스: 【】 제목 + 배경색 콜아웃
7) 빈 줄 단락 제거(HWP 변환 시 공백 단락 스킵)

권장 열람
- 서식·색이 더 잘 보이는 파일: DOCX (가독성판)
- 한글 제출용: HWP (가독성판)
- 한글에서 열은 뒤: 쪽 미리보기로 빈 쪽/큰 공백 확인 후, 필요 시 문단 간격만 추가 조정

크몽 반려 대응
- 이미지형 전자책이 아니라면 본문이 쪽을 충분히 채우도록 구성함
- 표지 이미지는 별도 디자인 시 여백 규격 준수
""",
        encoding="utf-8",
    )


def main():
    write_note()
    build()
    convert_hwp(DOCX, HWP)
    print("DONE", DIR_MS)
    for p in sorted(DIR_MS.glob("*")):
        print(p.name, p.stat().st_size)


if __name__ == "__main__":
    main()

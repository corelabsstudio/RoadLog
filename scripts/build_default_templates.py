"""로드로그 기본 제공 회사 서식 5종 생성 (xlsx + docx)."""
from __future__ import annotations

import json
from pathlib import Path

from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parent.parent
OUT = ROOT / "web" / "assets" / "templates"
OUT.mkdir(parents=True, exist_ok=True)

thin = Border(
    left=Side(style="thin", color="334155"),
    right=Side(style="thin", color="334155"),
    top=Side(style="thin", color="334155"),
    bottom=Side(style="thin", color="334155"),
)
header_fill = PatternFill("solid", fgColor="0F172A")
header_font = Font(bold=True, color="FFFFFF", name="맑은 고딕", size=10)
title_font = Font(bold=True, name="맑은 고딕", size=16, color="0F172A")
label_font = Font(bold=True, name="맑은 고딕", size=10)
cell_font = Font(name="맑은 고딕", size=10)
note_font = Font(name="맑은 고딕", size=9, color="64748B")


def make_xlsx(name: str, title: str, meta_fields: list, columns: list, rows: int = 8, note: str = "") -> Path:
    wb = Workbook()
    ws = wb.active
    ws.title = "운행일지"
    ncols = len(columns)

    ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=max(ncols, 4))
    ws["A1"] = title
    ws["A1"].font = title_font
    ws["A1"].alignment = Alignment(horizontal="center", vertical="center")
    ws.row_dimensions[1].height = 30

    r = 3
    for i, (lab, ph) in enumerate(meta_fields):
        col = 1 + (i % 2) * 2
        if i and i % 2 == 0:
            r += 1
        c_lab = ws.cell(r, col, lab)
        c_lab.font = label_font
        c_lab.border = thin
        c_lab.fill = PatternFill("solid", fgColor="F1F5F9")
        c_val = ws.cell(r, col + 1, ph)
        c_val.border = thin
        c_val.font = cell_font

    r += 2
    for c, colname in enumerate(columns, 1):
        cell = ws.cell(r, c, colname)
        cell.font = header_font
        cell.fill = header_fill
        cell.border = thin
        cell.alignment = Alignment(horizontal="center", vertical="center")

    for i in range(1, rows + 1):
        for c in range(1, ncols + 1):
            cell = ws.cell(r + i, c, "")
            cell.border = thin
            cell.font = cell_font
        ws.row_dimensions[r + i].height = 22

    sr = r + rows + 2
    ws.cell(sr, 1, "운행 요약").font = label_font
    ws.cell(sr, 1).border = thin
    ws.cell(sr, 1).fill = PatternFill("solid", fgColor="F1F5F9")
    ws.merge_cells(start_row=sr, start_column=2, end_row=sr + 1, end_column=ncols)
    for rr in (sr, sr + 1):
        for c in range(1, ncols + 1):
            ws.cell(rr, c).border = thin

    sg = sr + 3
    ws.cell(sg, max(1, ncols - 3), "작성자").font = label_font
    ws.cell(sg, max(1, ncols - 1), "확인자").font = label_font
    ws.cell(sg + 1, max(1, ncols - 3), "(서명)").font = cell_font
    ws.cell(sg + 1, max(1, ncols - 1), "(서명)").font = cell_font
    if note:
        ws.cell(sg + 3, 1, note).font = note_font

    for c in range(1, ncols + 1):
        ws.column_dimensions[get_column_letter(c)].width = 13 if c > 2 else 14

    path = OUT / f"{name}.xlsx"
    wb.save(path)
    return path


def set_run_font(run, size=10, bold=False):
    run.font.name = "맑은 고딕"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    run.font.size = Pt(size)
    run.bold = bold


def make_docx(name: str, title: str, meta_fields: list, columns: list, note: str = "") -> Path:
    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(1.5)
        s.bottom_margin = Cm(1.5)
        s.left_margin = Cm(1.5)
        s.right_margin = Cm(1.5)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    set_run_font(r, 16, True)

    n_meta_rows = (len(meta_fields) + 1) // 2
    mt = doc.add_table(rows=n_meta_rows, cols=4)
    mt.style = "Table Grid"
    for i, (lab, ph) in enumerate(meta_fields):
        row = i // 2
        col = (i % 2) * 2
        mt.rows[row].cells[col].text = lab
        mt.rows[row].cells[col + 1].text = ph or ""

    doc.add_paragraph("")
    tt = doc.add_table(rows=9, cols=len(columns))
    tt.style = "Table Grid"
    for c, colname in enumerate(columns):
        tt.rows[0].cells[c].text = colname

    doc.add_paragraph("")
    sp = doc.add_paragraph()
    r = sp.add_run("운행 요약: _______________________________________________")
    set_run_font(r, 10)
    sp2 = doc.add_paragraph()
    r = sp2.add_run("작성자 ______________      확인자 ______________")
    set_run_font(r, 10)
    if note:
        n = doc.add_paragraph()
        r = n.add_run(note)
        set_run_font(r, 9)

    path = OUT / f"{name}.docx"
    doc.save(path)
    return path


TEMPLATES = [
    {
        "name": "01_기본형_운행일지",
        "title": "차량 운행일지 (기본형)",
        "short": "기본형",
        "desc": "가장 많이 쓰는 일반 업무 운행 양식",
        "meta": [("작성일", ""), ("차량번호", ""), ("운전자", ""), ("부서", "")],
        "cols": ["순번", "출발시각", "도착시각", "출발지", "도착지", "운행목적", "거리(km)", "비고"],
        "note": "로드로그 기본 제공 서식 · 일반 업무 운행용",
    },
    {
        "name": "02_간단형_운행일지",
        "title": "차량 운행일지 (간단형)",
        "short": "간단형",
        "desc": "항목을 줄인 빠른 기록용 양식",
        "meta": [("작성일", ""), ("차량번호", ""), ("운전자", "")],
        "cols": ["시각", "출발지", "도착지", "목적", "거리(km)"],
        "note": "로드로그 기본 제공 서식 · 빠른 기록용",
    },
    {
        "name": "03_상세형_누적거리",
        "title": "차량 운행일지 (상세형 · 누적거리)",
        "short": "상세형",
        "desc": "누적거리·점심까지 적는 상세 양식",
        "meta": [
            ("작성일", ""),
            ("차량번호", ""),
            ("운전자", ""),
            ("회사명", ""),
            ("최초누적(km)", ""),
            ("종료누적(km)", ""),
        ],
        "cols": ["순번", "출발", "도착", "출발지", "도착지", "목적", "거리", "점심", "비고"],
        "note": "로드로그 기본 제공 서식 · 누적거리·점심 포함",
    },
    {
        "name": "04_영업외근형",
        "title": "차량 운행일지 (영업·외근형)",
        "short": "영업·외근형",
        "desc": "거래처 방문·영업 외근 기록용",
        "meta": [("작성일", ""), ("차량번호", ""), ("담당자", ""), ("거래처구분", "")],
        "cols": ["순번", "출발", "도착", "방문처", "담당자", "방문목적", "거리(km)", "결과/비고"],
        "note": "로드로그 기본 제공 서식 · 영업·외근 방문 기록",
    },
    {
        "name": "05_현장공사형",
        "title": "차량 운행일지 (현장·공사형)",
        "short": "현장·공사형",
        "desc": "현장 이동·공사 지원 운행용",
        "meta": [("작성일", ""), ("차량번호", ""), ("운전자", ""), ("현장명", "")],
        "cols": ["순번", "출발", "도착", "출발지", "도착지(현장)", "작업내용", "거리(km)", "동승/비고"],
        "note": "로드로그 기본 제공 서식 · 현장 이동·공사 지원용",
    },
]


def main() -> None:
    manifest = []
    for t in TEMPLATES:
        make_xlsx(t["name"], t["title"], t["meta"], t["cols"], note=t["note"])
        make_docx(t["name"], t["title"], t["meta"], t["cols"], note=t["note"])
        manifest.append(
            {
                "id": t["name"],
                "title": t["short"],
                "full_title": t["title"],
                "desc": t["desc"],
                "xlsx": f"/assets/templates/{t['name']}.xlsx",
                "docx": f"/assets/templates/{t['name']}.docx",
            }
        )
        print("ok", t["name"])

    (OUT / "manifest.json").write_text(
        json.dumps(manifest, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    print("manifest", len(manifest), "->", OUT)


if __name__ == "__main__":
    main()
